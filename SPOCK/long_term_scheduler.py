from alive_progress import alive_bar
from astropy.table import Table
from astropy import units as u
from astropy.coordinates import SkyCoord, get_sun, AltAz, EarthLocation
from astropy.utils import iers
from astropy.time import Time
from astropy.utils.data import clear_download_cache

clear_download_cache()
from astroplan import FixedTarget, AltitudeConstraint, MoonSeparationConstraint, AtNightConstraint, observability_table, \
    is_observable, months_observable, time_grid_from_range, LocalTimeConstraint, is_always_observable
from astroplan import TimeConstraint, Observer, moon_illumination
from colorama import Fore
from datetime import date, datetime, time, timedelta
from docx import Document
from docx.shared import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
from eScheduler.spe_schedule import SPECULOOSScheduler, Schedule, ObservingBlock, Transitioner
import functools
from functools import reduce
import gspread
import mphot
import numpy as np
import os
import os.path, time
from oauth2client.service_account import ServiceAccountCredentials
import pandas as pd
import paramiko
import requests
import shutil
import ssl #comment
import subprocess
import sys
import time
from tqdm.auto import tqdm
from SPOCK import user_portal, pwd_portal, pwd_appcs, pwd_SNO_Reduc1, user_chart_studio, \
    pwd_chart_studio, path_spock, target_list_from_stargate_path
import SPOCK.ETC as ETC

iers.IERS_A_URL = 'https://datacenter.iers.org/data/9/finals2000A.all'  # 'http://maia.usno.navy.mil/ser7/finals2000A.all'#'ftp://cddis.gsfc.nasa.gov/pub/products/iers/finals2000A.all'
ssl._create_default_https_context = ssl._create_unverified_context

from .make_night_plans import make_np, make_astra_schedule_file, offset_target_position
from .upload_night_plans import upload_np, upload_np_tn, upload_np_ts


def get_hours_files_sno(username='speculoos', password=pwd_SNO_Reduc1):
    """ get nb hours obs on SNO

    Returns
    -------
    txt file
        file ObservationHours.txt

    """
    hostname = '172.16.3.11'
    port = 22
    paramiko.util.log_to_file('paramiko.log')
    s = paramiko.SSHClient()
    s.load_system_host_keys()
    try:
        s.connect(hostname, port, username, password)
    except TimeoutError:
        sys.exit(Fore.RED + 'ERROR:  ' + Fore.BLACK + ' Make sure the VPN is connected')

    ftp_client = s.open_sftp()
    ftp_client.get('/home/speculoos/SNO/ObservationHours/ObservationHours.txt', path_spock +
                   '/survey_hours/ObservationHours.txt')
    ftp_client.close()

    s.close()
    df = pd.read_csv(path_spock + '/survey_hours/ObservationHours.txt', delimiter=',', skipinitialspace=True)
    df = df.sort_values(['Target'])
    df.to_csv(path_spock + '/survey_hours/ObservationHours.txt', sep=',', index=False)


def max_unit_list(x):
    """ return max of list

    Parameters
    ----------
    x : list

    Returns
    -------
    float
        max of list

    """
    a = max(x)
    return a.value


def first_elem_list(x):
    """ first element of list

    Parameters
    ----------
    x : list

    Returns
    -------
    float
        first element of list

    """
    a = x[0]
    return a


def last_elem_list(x):
    """ last element of list

    Parameters
    ----------
    x : list

    Returns
    -------
    float
        last element

    """
    a = x[-1]
    return a


def coord_transfotm_to_alt(x, frame):
    """ transform astropy.coordinates to a given frame

    Parameters
    ----------
    x  : astropy.coordinates
    coordinates
    frame : str
        frame in astropy.coordinates

    Returns
    -------
    astropy.coordinates
        coordinates in the frame chosen

    """
    a = x.coord.transform_to(frame).alt
    return a


def index_list1_list2(list1, list2):  # list 2 longer than list 1
    """ index of list1 in list2 and list2 in list1

    Parameters
    ----------
    list1 : list

    list2 : list

    Returns
    -------
    list
        list of index of list1 in list2 and list2 in list1

    """
    idx_list1_in_list2 = []
    idx_list2_in_list1 = []
    for i in range(len(list2)):
        for j in range(len(list1)):
            if list2[i] == list1[j]:
                idx_list1_in_list2.append(i)
                idx_list2_in_list1.append(j)
    return idx_list1_in_list2, idx_list2_in_list1


def diff_list(li1, li2):
    """ Inform on the difference between two lists

    Parameters
    ----------
    li1: list
    li2: list

    Returns
    -------
    list
        Elements than are in list 1 but not in list 2
    """

    return list(set(li1) - set(li2))


def compare_target_lists(path_target_list, user=user_portal, password=pwd_portal):
    """ Compare the target list from the given folder to the one on STARGATE and Cambridge server
    If different trigger a warning a tell how many targets are actually different from the referenced target list

    Parameters
    ----------
    path_target_list: str
        path on your computer toward the target list, by default take the one on the Cambridge server
    user: your portal username
    password: your portal password

    Returns
    -------
    print
         An idication about if the target list is the referenced one or not

    """
    targeturl = "http://www.mrao.cam.ac.uk/SPECULOOS/target_list_gaia.csv"
    resp = requests.get(targeturl, auth=(user, password))
    open('target_list_gaia.csv', 'wb').write(resp.content)
    content = resp.text.replace("\n", "")
    df_target_list_gaia = pd.read_csv('target_list_gaia.csv', delimiter=',', skipinitialspace=True,
                                      error_bad_lines=False)
    df_user = pd.read_csv(path_target_list, delimiter=' ', skipinitialspace=True, error_bad_lines=False)
    if diff_list(df_user['Name'], df_target_list_gaia['spc']):
        print(Fore.YELLOW + 'WARNING:  ' + Fore.BLACK +
              ' Targets in User\'s list but not in Cambridge server\'s list: ',
              diff_list(df_user['Name'], df_target_list_gaia['spc']))
    if diff_list(df_target_list_gaia['spc'], df_user['Name']):
        print(Fore.YELLOW + 'WARNING:  ' + Fore.BLACK +
              ' Targets in Cambridge server\'s list but not in User\'s list: ',
              diff_list(df_target_list_gaia['spc'], df_user['Name']))
    else:
        print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' OK ! User\'s list is similar to the one on the Cambridge server')


def sso_planned_targets(date_is, telescope):
    """ tell which target are scheduled on SSO on a given day

    Parameters
    ----------
    date_is: date of day in fmt 'yyyy-mm-dd'
    telescope : str
        name of telescope

    Returns
    -------
    list
        list of targets scheduled on this SSO telescope that  day

    """
    telescopes = ['Io', 'Europa', 'Ganymede', 'Callisto']
    if telescope in telescopes:
        telescopes.remove(telescope)
    targets_on_sso_telescopes = []

    #local
    for i in range(len(telescopes)):
        night_block_str = '/night_blocks_' + telescopes[i] + '_' + str(date_is) + '.txt'
        path = path_spock + '/DATABASE/' + telescopes[i] + '/Archive_night_blocks' + night_block_str
        try:
            c = pd.read_csv(path, delimiter=' ', index_col=False)
            for tar in c['target']:
                targets_on_sso_telescopes.append(tar)
        except FileNotFoundError:
            print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK + ' No plans in your local file for  ' +
                  telescopes[i] + ' on the ' + str(date_is) + ', looking online')
    #online
    for i in range(len(telescopes)):
        nightb_url = "https://speculoos.withastra.io/SPECULOOS/Observations/" + telescope + '/schedule/Archive_night_blocks/night_blocks_' + \
                     telescope + '_' + str(date_is) + '.txt'
        nightb = requests.get(nightb_url, auth=(user_portal, pwd_portal))

        if nightb.status_code == 404:
            sys.exit(Fore.RED + 'ERROR:  ' + Fore.BLACK + ' No plans on the server for this date')
        else:
            path_local = path_spock + '/DATABASE/' + telescope + '/Archive_night_blocks/night_blocks_' + telescope + '_' + \
                 str(date_is) + '.txt'
            open(path_local, 'wb').write(nightb.content)
            nb_online = pd.read_csv(path_local, delimiter=' ', index_col=False)     
            for tar in nb_online['target']:
                targets_on_sso_telescopes.append(tar)  
    
    targets_on_sso_telescopes = list(set(targets_on_sso_telescopes))
    return targets_on_sso_telescopes


def sno_planned_targets(date_is):
    """ tell which target are scheduled on SNO on a given day

    Parameters
    ----------
    date_is : date
        date of day in fmt 'yyyy-mm-dd'

    Returns
    -------
    list
        list of targets scheduled on this SNO that  day

    """
    telescopes = ['Artemis']
    targets_on_sno_telescopes = []

    #local
    for i in range(len(telescopes)): 
        night_block_str = '/night_blocks_' + telescopes[i] + '_' + str(date_is) + '.txt'
        path = path_spock + '/DATABASE/' + telescopes[i] + '/Archive_night_blocks' + night_block_str
        try:
            c = pd.read_csv(path, delimiter=' ', index_col=False)
            for tar in c['target']:
                targets_on_sno_telescopes.append(tar)
        except FileNotFoundError:
            print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK + ' No plans in your local file for  ' +
                  telescopes[i] + ' on the ' + str(date_is) + ', looking online')
            
    #online
    for i in range(len(telescopes)):
        nightb_url = "https://speculoos.withastra.io/SPECULOOS/Observations/" + telescopes[i] + '/schedule/Archive_night_blocks/night_blocks_' + telescopes[i] + '_' + str(date_is) + '.txt'
        nightb = requests.get(nightb_url, auth=(user_portal, pwd_portal))

        if nightb.status_code == 404:
            sys.exit(Fore.RED + 'ERROR:  ' + Fore.BLACK + ' No plans on the server for this date')
        else:
            path_local = path_spock + '/DATABASE/' + telescopes[i] + '/Archive_night_blocks/night_blocks_' + telescopes[i] + '_' + \
                 str(date_is) + '.txt'
            open(path_local, 'wb').write(nightb.content)
            nb_online = pd.read_csv(path_local, delimiter=' ', index_col=False)     
            for tar in nb_online['target']:
                targets_on_sno_telescopes.append(tar)  
    
    targets_on_sno_telescopes = list(set(targets_on_sno_telescopes))
    return targets_on_sno_telescopes

def saintex_planned_targets(date_is):
    """ tell which target are scheduled on SNO on a given day
    Parameters
    ----------
    date_is : date
        date of day in fmt 'yyyy-mm-dd'

    Returns
    -------
    list
        list of targets scheduled on this Saint-Ex that  day

    """
    telescopes = ['Saint-Ex']
    targets_on_saintex_telescopes = []

    #local
    for i in range(len(telescopes)): 
        night_block_str = '/night_blocks_' + telescopes[i] + '_' + str(date_is) + '.txt'
        path = path_spock + '/DATABASE/' + telescopes[i] + '/Archive_night_blocks' + night_block_str
        try:
            c = pd.read_csv(path, delimiter=' ', index_col=False)
            for tar in c['target']:
                targets_on_saintex_telescopes.append(tar)
        except FileNotFoundError:
            print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK + ' No plans in your local file for  ' +
                  telescopes[i] + ' on the ' + str(date_is) + ', looking online')
            
    #online
    for i in range(len(telescopes)):
        nightb_url = "https://speculoos.withastra.io/SPECULOOS/Observations/" + telescopes[i] + '/schedule/Archive_night_blocks/night_blocks_' + telescopes[i] + '_' + str(date_is) + '.txt'
        nightb = requests.get(nightb_url, auth=(user_portal, pwd_portal))

        if nightb.status_code == 404:
            sys.exit(Fore.RED + 'ERROR:  ' + Fore.BLACK + ' No plans on the server for this date')
        else:
            path_local = path_spock + '/DATABASE/' + telescopes[i] + '/Archive_night_blocks/night_blocks_' + telescopes[i] + '_' + \
                 str(date_is) + '.txt'
            open(path_local, 'wb').write(nightb.content)
            nb_online = pd.read_csv(path_local, delimiter=' ', index_col=False)     
            for tar in nb_online['target']:
                targets_on_saintex_telescopes.append(tar)  
    
    targets_on_saintex_telescopes = list(set(targets_on_saintex_telescopes))
    return targets_on_saintex_telescopes

# def ts_planned_targets(date_is):
#     """ tell which target are scheduled on TS on a given day

#     Parameters
#     ----------
#     date_is : date
#         date of day in fmt 'yyyy-mm-dd'

#     Returns
#     -------
#     list
#         list of targets scheduled on this TS that  day

#     """
#     telescopes = ['TS_La_Silla']
#     targets_on_ts_telescopes = []
#     for i in range(len(telescopes)):
#         night_block_str = '/night_blocks_' + telescopes[i] + '_' + str(date_is) + '.txt'
#         path = path_spock + '/DATABASE/' + telescopes[i] + '/Archive_night_blocks/' + night_block_str
#         try:
#             c = pd.read_csv(path, delimiter=' ', index_col=False)
#             for tar in c['target']:
#                 targets_on_ts_telescopes.append(tar)
#         except FileNotFoundError:
#             print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK + ' No plans in your local file for  ' + telescopes[i] +
#                   ' on the ' + str(date_is))
#     return targets_on_ts_telescopes


# def tn_planned_targets(date_is):
#     """ tell which target are scheduled on TN on a given day

#     Parameters
#     ----------
#     date_is : date
#         date of day in fmt 'yyyy-mm-dd'

#     Returns
#     -------
#     list
#         list of targets scheduled on this TN that  day

#     """
#     telescopes = ['TN_Oukaimeden']
#     targets_on_tn_telescopes = []
#     for i in range(len(telescopes)):
#         night_block_str = '/night_blocks_' + telescopes[i] + '_' + str(date_is) + '.txt'
#         path = path_spock + '/DATABASE/' + telescopes[i] + '/Archive_night_blocks/' + night_block_str
#         try:
#             c = pd.read_csv(path, delimiter=' ', index_col=False)
#             for tar in c['target']:
#                 targets_on_tn_telescopes.append(tar)
#         except FileNotFoundError:
#             print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK + ' No plans in your local file for  ' + telescopes[i] +
#                   ' on the ' + str(date_is))
#     return targets_on_tn_telescopes


def target_list_good_coord_format(path_target_list):
    """ Give target corrdinates in ICRS format (used for astropy.coordinates SkyCoord function)

    Parameters
    ----------
    path_target_list: str
        path on your computer toward the target list, by default take the one on the Cambridge server

    Returns
    -------
    targets: astropy.FixedTarget
        targets list with the following format : [<FixedTarget "Sp0002+0115" at SkyCoord (ICRS):
        (ra, dec) in deg (0.52591667, 1.26003889)>,


    """
    df = pd.read_csv(path_target_list, delimiter=',')
    target_table_spc = Table.from_pandas(df)
    targets = [FixedTarget(coord=SkyCoord(ra=target_table_spc['RA'][i] * u.degree,
                                          dec=target_table_spc['DEC'][i] * u.degree),
                           name=target_table_spc['Sp_ID'][i]) for i in range(len(target_table_spc['RA']))]
    return targets


def charge_observatories(Name):
    """

    Parameters
    ----------
    Name : str
        name of the observatory (ex: 'SSO')

    Returns
    -------
    astroplan.observer
        all info on observatory loaded

    """
    observatories = []
    # Observatories
    if 'SSO' in str(Name):
        location = EarthLocation.from_geodetic(-70.40300000000002 * u.deg, -24.625199999999996 * u.deg,
                                               2635.0000000009704 * u.m)
        observatories.append(Observer(location=location, name="SSO", timezone="UTC"))

    if 'SNO' in str(Name):
        location_SNO = EarthLocation.from_geodetic(-16.50583131 * u.deg, 28.2999988 * u.deg, 2390 * u.m)
        observatories.append(Observer(location=location_SNO, name="SNO", timezone="UTC"))

    if 'Saint-Ex' in str(Name):
        location_saintex = EarthLocation.from_geodetic(-115.48694444444445 * u.deg, 31.029166666666665 * u.deg,
                                                       2829.9999999997976 * u.m)
        observatories.append(Observer(location=location_saintex, name="Saint-Ex", timezone="UTC"))

    if 'TS_La_Silla' in str(Name):
        location_TSlasilla = EarthLocation.from_geodetic(-70.73000000000002 * u.deg, -29.25666666666666 * u.deg,
                                                         2346.9999999988418 * u.m)
        observatories.append(Observer(location=location_TSlasilla, name="TS_La_Silla", timezone="UTC"))

    if 'TN_Oukaimeden' in str(Name):
        location_TNOuka = EarthLocation.from_geodetic(-7.862263 * u.deg, 31.20516 * u.deg, 2751 * u.m)
        observatories.append(Observer(location=location_TNOuka, name="TN_Oukaimeden", timezone="UTC"))

    if 'Munich' in str(Name):
        location_munich = EarthLocation.from_geodetic(48.2 * u.deg, -11.6 * u.deg, 600 * u.m)
        observatories.append(Observer(location=location_munich, name="Munich", timezone="UTC"))

    return observatories


def _generate_24hr_grid(t0, start, end, N, for_deriv=False):
    """ Generate a nearly linearly spaced grid of time durations.
    The midpoints of these grid points will span times from ``t0``+``start``
    to ``t0``+``end``, including the end points, which is useful when taking
    numerical derivatives.
    Parameters
    ----------
    t0 : `~astropy.time.Time`
        Time queried for, grid will be built from or up to this time.
    start : float
        Number of days before/after ``t0`` to start the grid.
    end : float
        Number of days before/after ``t0`` to end the grid.
    N : int
        Number of grid points to generate
    for_deriv : bool
        Generate time series for taking numerical derivative (modify
        bounds)?
    Returns
    -------
    `~astropy.time.Time`
    """

    if for_deriv:
        time_grid = np.concatenate([[start - 1 / (N - 1)],
                                    np.linspace(start, end, N)[1:-1],
                                    [end + 1 / (N - 1)]]) * u.day
    else:
        time_grid = np.linspace(start, end, N) * u.day

    # broadcast so grid is first index, and remaining shape of t0
    # falls in later indices. e.g. if t0 is shape (10), time_grid
    # will be shape (N, 10). If t0 is shape (5, 2), time_grid is (N, 5, 2)
    while time_grid.ndim <= t0.ndim:
        time_grid = time_grid[:, np.newaxis]
    # we want to avoid 1D grids since we always want to broadcast against targets
    if time_grid.ndim == 1:
        time_grid = time_grid[:, np.newaxis]
    return t0 + time_grid


def altaz(self, time, target=None, obswl=None, grid_times_targets=False):
    """

    Parameters
    ----------
    self
    time
    target
    obswl
    grid_times_targets

    Returns
    -------

    """

    if target is not None:
        time, target = self._preprocess_inputs(time, target, grid_times_targets)

        altaz_frame = AltAz(location=self.location, obstime=time,
                            pressure=self.pressure, obswl=obswl,
                            temperature=self.temperature,
                            relative_humidity=self.relative_humidity)
    if target is None:
        # Return just the frame
        return altaz_frame
    else:
        return target.transform_to(altaz_frame)


def observability(j, time_range, observatory, targets, constraints):
    """ Give a table with the observability score for each target of targets
    regarding the constraints given and for all ranges of time_range

    Parameters
    ----------
        j : list
            [start end], element of time range (that is to say, month of the year)
        time_range : list
            of astropy.Time range with start and end times
        observatory : astroplan.observer
            observatory chosen
        targets : astropy
            target list on the FixedTarget() format from astroplan
        constraints : astroplan.constraint
            general constraints for a target to be shceduled

    Returns
    -------
        Observability table: astropy.table.Table
             12 columns (target name and each element of time_range, e.g months),
        rows= nb of targets
    """
    targets_observable = []
    observable = np.argwhere(is_observable(constraints, observatory, targets, time_range=time_range))

    # WARNING: Need to be replace by a np.where, but I can't find how
    [targets_observable.append(targets[int(obs)]) for obs in observable]

    table_observability = observability_table(constraints, observatory, targets_observable, time_range=time_range,
                                              time_grid_resolution=0.5 * u.hour)  # calcul une deuxieme fois is observable
    table_observability.remove_column('always observable')
    table_observability.remove_column('ever observable')
    table_observability.rename_column('fraction of time observable', 'Month' + str(j))

    return table_observability


def reverse_observability(observatory, targets, constraints, time_ranges):
    """
    Reverse observability table, rows become columns
    Parameters
    ----------
    observatory : str
        observatory chosen
    targets : list of astropy.FixedTarget
        target list on the FixedTarget() format from astroplan
    constraints : astroplan.constraints
        general constraints for a target to be scheduled
    time_ranges : list
        List of astropy.Time range with start and end times

    Returns
    -------
    reverse_df1 : astropy.table
        observability table with no NaN value (0 instead) inversed with targets as columns
        and elements of time_ranges (months) as rows
    """
    start_fmt = Time(time_ranges[0][0].iso, out_subfmt='date').iso
    end_fmt = Time(time_ranges[len(time_ranges) - 1][1].iso, out_subfmt='date').iso

    if os.path.exists(path_spock + '/SPOCK_files/reverse_Obs_' + str(observatory.name) + '_' + start_fmt + '_' +
                      end_fmt + '_' + str(len(targets)) + '.csv'):
        name_file = path_spock + '/SPOCK_files/reverse_Obs_' + str(observatory.name) + '_' + start_fmt + '_' + \
                    end_fmt + '_' + str(len(targets)) + '.csv'
        reverse_df1 = pd.read_csv(name_file, delimiter=',')
        return reverse_df1

    else:
        tables_observability = list(map((lambda x: observability(x, time_ranges[x], observatory, targets, constraints)),
                                        range(0, len(time_ranges))))
        df = list(map((lambda x: pd.DataFrame(tables_observability[x].to_pandas())), range(0, len(time_ranges))))
        a = reduce(functools.partial(pd.merge, how='outer', on='target name'), df)
        df = a.replace(to_replace=float('NaN'), value=0.0)
        df1 = df.set_index('target name')
        reverse_df1 = df1.T
        reverse_df1.to_csv(path_spock + '/SPOCK_files/reverse_Obs_' + str(observatory.name) + '_' + start_fmt + '_' +
                           end_fmt + '_' + str(len(targets)) + '.csv', sep=',')
        return reverse_df1


def month_option(target_name, reverse_df1):
    """ create a list of the best month for oservation for each target

    Parameters
    ----------
    target_name : str
        name of the target
    reverse_df1 : astropy.table
        observability table with no NaN value (0 instead) inversed with targets as columns
        and elements of time_ranges (months) as rows

    Returns
    -------
    month : list
        a list with the best month to observe the target
    month_2nd_option : list
        same but for the second best month
    months_3rd_option : list
        same but for the third best month
    months_4th_option : list
        same but for the fourth best month
    months_5th_option : list
        same but for the fiveth best month


    Remarks
    -------
        the 2nd, 3rd etc choices are here in case the target list is not long enough to give
        solutions for all the telescopes each months, allows to avoid blancks in observations
    """

    try:
        months = reverse_df1[str(target_name)].idxmax()
        months_2nd_option = reverse_df1[str(target_name)].nlargest(2, keep='first').index[1]
        months_3rd_option = reverse_df1[str(target_name)].nlargest(3, keep='first').index[2]
        months_4th_option = reverse_df1[str(target_name)].nlargest(4, keep='first').index[3]
        months_5th_option = reverse_df1[str(target_name)].nlargest(5, keep='first').index[4]

    except KeyError:
        months = 0
        months_2nd_option = 0
        months_3rd_option = 0
        months_4th_option = 0
        months_5th_option = 0
    return [months, months_2nd_option, months_3rd_option, months_4th_option, months_5th_option]


def save_schedule(save, over_write, date_range, telescope):
    """ save schedules in destination

    Parameters
    ----------
    save : bool
        True if want to save, False if not
    over_write : bool
        True if want overwrite, False if not
    date_range : list of date
         list with 2 elements, start date  and end date
    telescope : str
        name of telescope for which you which to save the schedules

    Returns
    -------
    message

    """

    date_range_in_days = int((date_range[1] - date_range[0]).value)
    for i in range(0, date_range_in_days):
        day = date_range[0] + i
        if save:
            source = path_spock + '/' + 'night_blocks_propositions/' + 'night_blocks_' + telescope + '_' + \
                     day.tt.datetime.strftime("%Y-%m-%d") + '.txt'
            destination = path_spock + '/DATABASE/' + telescope + '/'
            destination_2 = path_spock + '/DATABASE/' + telescope + '/' + 'Archive_night_blocks/'
            if over_write:
                # dest = shutil.copy(source, destination)
                dest2 = shutil.copy(source, destination_2)
                # print(Fore.GREEN + 'INFO:  ' + Fore.BLACK + '\"' + source + '\"' + ' has been over-written to ' +
                #       '\"' + destination + '\"' )
                print(Fore.GREEN + 'INFO:  ' + Fore.BLACK + '\"' + source + '\"' + ' has been copied to ' + '\"' +
                      destination_2 + '\"')
            if not over_write:
                try:
                    dest = shutil.move(source, destination)
                    print(Fore.GREEN + 'INFO:  ' + Fore.BLACK + '\"' + source + '\"' + ' has been copied to ' +
                          '\"' + destination + '\"')
                except shutil.Error:
                    print(Fore.GREEN + 'INFO:  ' + Fore.BLACK + '\"' + destination + 'night_blocks_' + telescope +
                          '_' + day.tt.datetime.strftime("%Y-%m-%d") + '.txt' + '\"' + ' already exists')
        if not save:
            print(Fore.GREEN + 'INFO:  ' + Fore.BLACK + ' Those plans have not been saved')


def make_plans(day, nb_days, telescope):
    """ make plans for telescope for a certain number of day from a start day

    Parameters
    ----------
    day : int
        day
    nb_days : int
        number of days
    telescope : string
        name telescope

    Returns
    -------


    """

    make_np(day, nb_days, telescope)
    make_astra_schedule_file(day, nb_days, telescope)


def upload_plans(day, nb_days, telescope):
    """ upload plans to DATABASE

    Parameters
    ----------
    day : date
        date in fmt 'yyyy-mm-dd'
    nb_days : int
        number of days
    telescope : str
        name of telescope

    Returns
    -------

    """
    if (telescope == 'Io') or telescope == ('Europa') or (telescope == 'Ganymede') or (telescope == 'Callisto') or (telescope =='Artemis')\
            or (telescope =='Saint-Ex'):
        upload_np(day,nb_days,telescope)
    #if telescope.find('TS_La_Silla') is not -1:
    #    upload_np_ts(day, nb_days)
    #if telescope.find('TN_Oukaimeden') is not -1:
    #    upload_np_tn(day, nb_days)

    # ------------------- update archive date by date plans folder  ------------------

    # path_database = os.path.join('speculoos@appcs.ra.phy.cam.ac.uk:/appct/data/SPECULOOSPipeline/',
    #                              telescope, 'schedule')
    path_gant_chart = os.path.join(path_spock + '/SPOCK_Figures/Preview_schedule.html')
    path_gant_chart_masterfile = os.path.join('/Users/elsaducrot/spock_2/SPOCK_Files/spock_stats_masterfile.csv')
    path_database_home = \
        os.path.join('speculoos@appcs.ra.phy.cam.ac.uk:/appct/data/SPECULOOSPipeline/spock_files/Preview_schedule.html')
    path_database_home_masterfile = \
        os.path.join('speculoos@appcs.ra.phy.cam.ac.uk:/appct/data/SPECULOOSPipeline/spock_files/'
                     'spock_stats_masterfile.csv')
    # print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' Path local \'Gant chart\' = ', path_gant_chart)
    # print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' Path database \'Gant chart\' = ',  path_database_home)
    # subprocess.Popen(["sshpass", "-p", pwd_appcs, "scp", "-r", path_gant_chart, path_database_home])
    # subprocess.Popen(["sshpass", "-p", pwd_appcs, "scp", "-r", path_gant_chart_masterfile,
    #                   path_database_home_masterfile])


class Schedules:
    """
    Class to Make schedules for the target list, observatory, date_range and startegy indicated

    """

    def __init__(self):

        self.Altitude_constraint = 25
        self.constraints = None
        self.date_range = None  # date_range
        self.dur_obs_set_target = None
        self.dur_obs_rise_target = None
        self.dur_obs_both_target = None
        self.duration_segments = None  # duration_segments
        self.first_target = None
        self.first_target_by_day = []
        self.idx_first_target = None
        self.idx_first_target_by_day = []
        self.idx_second_target = None
        self.idx_second_target_by_day = []
        self.index_prio = None
        self.index_prio_by_day = []
        self.idx_planned_sso = None
        self.idx_SSO_in_planned = None
        self.Moon_constraint = 30
        self.moon_and_visibility_constraint_table = None
        self.nb_segments = None  # nb_segments
        self.night_block = []
        self.night_block_by_day = []
        self.observatory = None  # observatory
        self.observability_table_day = None
        self.planned_targets = []
        self.priority = None
        self.priority_by_day = []
        self.priority_ranked = None
        self.priority_ranked_by_day = []
        self.read_locked_target = True
        self.reverse_df1 = None
        self.second_target = None
        self.second_target_by_day = []
        self.strategy = None
        self.targets = None
        self.target_list = None
        self.target_table_spc = []
        self.telescopes = ['Io', 'Europa', 'Ganymede', 'Callisto', 'TS_La_Silla', 'TN_Oukaimeden']
        self.telescope = []
        #self.start_between_civil_nautical = None
        #self.end_between_civil_nautical = None
        self.start_night = None
        self.end_night = None
        self.night_duration = None

    @property
    def idx_rise_targets_sorted(self):
        """ index for rise targets sorted
        
        Returns
        -------
        list of int
            index of rise targets sorted

        """
        idx_rise_targets = (self.priority_ranked['set or rise'] == 'rise')
        idx_rise_targets_sorted = self.index_prio[idx_rise_targets]
        return idx_rise_targets_sorted

    @property
    def idx_set_targets_sorted(self):
        """ index for set targets sorted

        Returns
        -------
        list of int
            index of set targets sorted
        """
        idx_set_targets = (self.priority_ranked['set or rise'] == 'set')
        idx_set_targets_sorted = self.index_prio[idx_set_targets]
        return idx_set_targets_sorted

    @property
    def months_obs(self):
        """ month of obs

        Returns
        -------
        int
            month number (between 0 and 11)

        """
        date_format = "%Y-%m-%d %H:%M:%S.%f"
        for i, t in enumerate(self.time_ranges):
            if (datetime.strptime(t[0].value, date_format) <=
                datetime.strptime(self.date_range[0].value, date_format) <=
                datetime.strptime(t[1].value, date_format)) \
                    and (datetime.strptime(t[0].value, date_format) <=
                         datetime.strptime(self.date_range[1].value, date_format) <=
                         datetime.strptime(t[1].value, date_format)):
                return i
            if (datetime.strptime(t[0].value, date_format) <=
                datetime.strptime(self.date_range[0].value, date_format) <= datetime.strptime(t[1].value, date_format)) \
                    and (datetime.strptime(t[1].value, date_format) <=
                         datetime.strptime(self.date_range[1].value, date_format)):
                if i < (len(self.time_ranges) - 1):
                    return i + 1
                if i == (len(self.time_ranges) - 1):
                    return i

    @property
    def date_range_in_days(self):
        """ number of days in date range

        Returns
        -------
        int
            number of day between date start and date end

        """
        date_format = "%Y-%m-%d %H:%M:%S.%f"
        date_start = datetime.strptime(self.date_range[0].value, date_format)
        date_end = datetime.strptime(self.date_range[1].value, date_format)
        return (date_end - date_start).days

    @property
    def nb_hours_threshold(self):
        """ number of hours to reach

        Returns
        -------
        list
            list as long as the target list

        """
        nb_hours_threshold = [100] * len(self.target_table_spc)
        return nb_hours_threshold

    @property
    def date_ranges_day_by_day(self):
        """ date range day by day

        Returns
        -------
        list
            list of date ranges

        """
        date_format = "%Y-%m-%d %H:%M:%S.%f"
        d2 = datetime.strptime(self.date_range[1].value, date_format)
        i = 0
        t = datetime.strptime(self.date_range[0].value, date_format)
        u = t
        date_ranges_day_by_day = []
        while t < d2:
            d = timedelta(days=1)
            t = u + d * i
            date_ranges_day_by_day.append(Time(t))
            i += 1
        return date_ranges_day_by_day

    @property
    def nb_hours_observed(self):
        """ nb hours observed for each target

        Returns
        -------
        list
            nb hours observed

        """
        nb_hours_observed = self.target_table_spc['nb_hours_surved']
        return nb_hours_observed

    @property
    def time_ranges(self):
        year = self.date_range[0].tt.datetime.strftime("%Y")
        time_ranges = [Time([year + '-01-01 12:00:00', year + '-01-31 12:00:00']),
                       Time([year + '-02-01 12:00:00', year + '-02-28 12:00:00']),
                       Time([year + '-03-01 15:00:00', year + '-03-31 15:00:00']),
                       Time([year + '-04-01 15:00:00', year + '-04-30 15:00:00']),
                       Time([year + '-05-01 15:00:00', year + '-05-31 15:00:00']),
                       Time([year + '-06-01 15:00:00', year + '-06-30 15:00:00']),
                       Time([year + '-07-01 12:00:00', year + '-07-31 12:00:00']),
                       Time([year + '-08-01 12:00:00', year + '-08-31 12:00:00']),
                       Time([year + '-09-01 12:00:00', year + '-09-30 12:00:00']),
                       Time([year + '-10-01 12:00:00', year + '-10-31 12:00:00']),
                       Time([year + '-11-01 12:00:00', year + '-11-30 12:00:00']),
                       Time([year + '-12-01 12:00:00', year + '-12-31 12:00:00'])]
        return time_ranges

    def idx_SSO_observed_targets(self):
        df_cambridge = pd.read_csv(path_spock + '/survey_hours/SurveyTotal.txt', delimiter=' ', skipinitialspace=True,
                                   error_bad_lines=False)
        df_cambridge['Target'] = [x.strip().replace('SP', 'Sp') for x in df_cambridge['Target']]
        server_in_targetlist, targetlist_in_server = index_list1_list2(df_cambridge['Target'],
                                                                       self.target_table_spc['Sp_ID'])
        return server_in_targetlist

    def idx_SNO_observed_targets(self):
        df_artemis = pd.read_csv(path_spock + '/survey_hours/ObservationHours.txt', delimiter=',')
        artemis_in_targetlist, targetlist_in_artemis = index_list1_list2(df_artemis['Target'],
                                                                         self.target_table_spc['Sp_ID'])
        return artemis_in_targetlist

    def idx_saintex_observed_targets(self):
        df_saintex = pd.read_csv(path_spock + '/survey_hours/ObservationHours_Saint-Ex.txt', delimiter=',')
        saintex_in_targetlist, targetlist_in_saintex = index_list1_list2(df_saintex['Target'],
                                                                         self.target_table_spc['Sp_ID'])
        return saintex_in_targetlist

    def idx_trappist_observed_targets(self):
        df_trappist = pd.read_csv(path_spock + '/survey_hours/ObservationHours_TRAPPIST.txt', delimiter=',')
        trappist_in_targetlist, targetlist_in_trappist = index_list1_list2(df_trappist['Target'],
                                                                           self.target_table_spc['Sp_ID'])
        return trappist_in_targetlist

    def load_parameters(self, date_range=None):
        self.observatory = charge_observatories(self.observatory_name)[0]
        if date_range is not None:
            self.date_range = Time(date_range)
        self.strategy = 'continuous'
        # self.target_list = path_spock + '/target_lists/speculoos_target_list_v6.txt'
        self.target_list = target_list_from_stargate_path
        self.constraints = [AtNightConstraint()]
        df = pd.read_csv(self.target_list, delimiter=',')
        self.target_table_spc = Table.from_pandas(df)
        self.targets = target_list_good_coord_format(self.target_list)

        # --- UNCOMMENT IF YOU USE LOCAL FILE FOR TARGET LIST (NOT STARGATE) ---
        # last_mod = time.ctime(os.path.getmtime(self.target_list))
        # now = datetime.now()
        # current_time = now.strftime("%Y-%m-%d %H:%M:%S")
        # time_since_last_update = (Time(datetime.strptime(last_mod, "%a %b %d %H:%M:%S %Y"), format='datetime') - \
        #                           Time(datetime.strptime(current_time, "%Y-%m-%d %H:%M:%S"),
        #                                format='datetime')).value * 24
        # # self.update_nb_hours_all()
        # if abs(time_since_last_update) > 24:  # in hours
        #     print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' Updating the number of hours observed')
        #     self.update_nb_hours_all()
        # if self.date_range[1] <= self.date_range[0]:
        #     sys.exit(Fore.RED + 'ERROR:  ' + Fore.BLACK + ' end date inferior to start date')

    def update_nb_hours_all(self, user=user_portal, password=pwd_portal):
        # *********** TRAPPIST ***********
        # self.get_hours_files_trappist()

        # *********** SSO & SNO ***********
        self.update_telescope_from_server()  # get hours SSO
        TargetURL = "http://www.mrao.cam.ac.uk/SPECULOOS/reports/SurveyTotal"
        target_list = pd.read_csv(self.target_list, delimiter=',')
        target_list['telescope'] = ['None'] * len(target_list['telescope'])
        resp = requests.get(TargetURL, auth=(user, password))
        content = resp.text.replace("\n", "")
        open(path_spock + '/survey_hours/SurveyTotal.txt', 'wb').write(resp.content)
        df = pd.read_csv(path_spock + '/survey_hours/SurveyTotal.txt', delimiter=' ', skipinitialspace=True,
                         error_bad_lines=False)
        df = df.sort_values(['Target'])
        df.to_csv(path_spock + '/survey_hours/SurveyTotal.txt', sep=' ', index=False)
        df_camserver = pd.read_csv(path_spock + '/survey_hours/SurveyTotal.txt', delimiter=' ', skipinitialspace=True,
                                   error_bad_lines=False)
        df_camserver['Target'] = [x.strip().replace('SP', 'Sp') for x in df_camserver['Target']]

        # *********** Saint-Ex ***********
        df = pd.read_csv(path_spock + '/survey_hours/ObservationHours_Saint-Ex.txt', delimiter=',')
        df = df.sort_values(['Target'])
        df.to_csv(path_spock + '/survey_hours/ObservationHours_Saint-Ex.txt', sep=',', index=False)

        # Read files
        df_saintex = pd.read_csv(path_spock + '/survey_hours/ObservationHours_Saint-Ex.txt', delimiter=',')
        df_trappist = pd.read_csv(path_spock + '/survey_hours/ObservationHours_TRAPPIST.txt', delimiter=',')
        df_camserver_telescope = pd.read_csv(path_spock + '/survey_hours/SurveyByTelescope.txt', delimiter=' ',
                                             skipinitialspace=True, error_bad_lines=False)
        df_camserver_telescope['Target'] = [x.strip().replace('SP', 'Sp') for x in df_camserver_telescope['Target']]
        for i in range(len(target_list)):
            idxs = np.where((df_camserver_telescope['Target'] == target_list['Sp_ID'][i]))[0]
            target_list['telescope'][i] = list(df_camserver_telescope['Telescope'][idxs])

        saintex_in_targetlist, targetlist_in_saintex = index_list1_list2(df_saintex['Target'], target_list['Sp_ID'])
        camserver_in_targetlist, targetlist_in_camserver = index_list1_list2(df_camserver['Target'],
                                                                             target_list['Sp_ID'])
        trappist_in_targetlist, targetlist_in_trappist = index_list1_list2(df_trappist['Target'], target_list['Sp_ID'])

        with alive_bar(len(target_list)) as bar:
            for i in range(len(target_list)):
                bar()
                time.sleep(0.001)

                if np.any((np.asarray(camserver_in_targetlist) == i)):  # SSO or SNO
                    idx_camserver = np.argwhere((np.asanyarray(camserver_in_targetlist) == i))[0][0]
                    target_list['nb_hours_surved'][i] = df_camserver['Hours'][targetlist_in_camserver[idx_camserver]]

                    if np.any((np.asarray(saintex_in_targetlist) == i)) \
                            and np.any((np.asarray(trappist_in_targetlist) == i)):  # Saint-Ex and TRAPPIST
                        idx_saintex = np.argwhere((np.asanyarray(saintex_in_targetlist) == i))[0][0]
                        target_list['nb_hours_surved'][i] += \
                            df_saintex[' Observation Hours '][targetlist_in_saintex[idx_saintex]]
                        target_list['telescope'][i].append('Saint-Ex')
                        idx_trappist = np.argwhere((np.asanyarray(trappist_in_targetlist) == i))[0][0]
                        target_list['nb_hours_surved'][i] += \
                            df_trappist[' Observation Hours '][targetlist_in_trappist[idx_trappist]]
                        target_list['telescope'][i].append('TRAPPIST')

                    if np.any((np.asarray(trappist_in_targetlist) == i)) \
                            and np.all((np.asarray(saintex_in_targetlist) != i)):  # TRAPPIST but not Saint-Ex
                        idx_trappist = np.argwhere((np.asanyarray(trappist_in_targetlist) == i))[0][0]
                        target_list['nb_hours_surved'][i] += \
                            df_trappist[' Observation Hours '][targetlist_in_trappist[idx_trappist]]
                        target_list['telescope'][i].append('TRAPPIST')

                    if np.any((np.asarray(saintex_in_targetlist) == i)) \
                            and np.all((np.asarray(trappist_in_targetlist) != i)):  # Saint-Ex but not TRAPPIST
                        idx_saintex = np.argwhere((np.asanyarray(saintex_in_targetlist) == i))[0][0]
                        target_list['nb_hours_surved'][i] += \
                            df_saintex[' Observation Hours '][targetlist_in_saintex[idx_saintex]]
                        target_list['telescope'][i].append('Saint-Ex')

                if np.all((np.asarray(camserver_in_targetlist) != i)):  # Neither SSO nor SNO

                    if np.any((np.asarray(saintex_in_targetlist) == i)) \
                            and np.any((np.asarray(trappist_in_targetlist) == i)):  # Saint-Ex and TRAPPIST
                        idx_saintex = np.argwhere((np.asanyarray(saintex_in_targetlist) == i))[0][0]
                        target_list['nb_hours_surved'][i] = \
                            df_saintex[' Observation Hours '][targetlist_in_saintex[idx_saintex]]
                        target_list['telescope'][i].append('Saint-Ex')
                        idx_trappist = np.argwhere((np.asanyarray(trappist_in_targetlist) == i))[0][0]
                        target_list['nb_hours_surved'][i] += \
                            df_trappist[' Observation Hours '][targetlist_in_trappist[idx_trappist]]
                        target_list['telescope'][i].append('TRAPPIST')

                    if np.any((np.asarray(trappist_in_targetlist) == i)) \
                            and np.all((np.asarray(saintex_in_targetlist) != i)):  # TRAPPIST but not Saint-Ex
                        idx_trappist = np.argwhere((np.asanyarray(trappist_in_targetlist) == i))[0][0]
                        target_list['nb_hours_surved'][i] = \
                            df_trappist[' Observation Hours '][targetlist_in_trappist[idx_trappist]]
                        target_list['telescope'][i].append('TRAPPIST')

                    if np.any((np.asarray(saintex_in_targetlist) == i)) \
                            and np.all((np.asarray(trappist_in_targetlist) != i)):  # Saint-Ex but not TRAPPIST
                        idx_saintex = np.argwhere((np.asanyarray(saintex_in_targetlist) == i))[0][0]
                        target_list['nb_hours_surved'][i] = \
                            df_saintex[' Observation Hours '][targetlist_in_saintex[idx_saintex]]
                        target_list['telescope'][i].append('Saint-Ex')

        target_list.to_csv(self.target_list, sep=' ', index=False)
        target_list.to_csv(path_spock + '/target_lists/speculoos_target_list_v6_sep_coma.csv', sep=',')
        # subprocess.Popen(["sshpass", "-p", pwd_appcs, "scp", path_spock + '/target_lists/speculoos_target_list_v6.txt',
        #                   'speculoos@appcs.ra.phy.cam.ac.uk:/appct/data/SPECULOOSPipeline/spock_files/target_lists/'])

    def get_hours_files_trappist(self):
        scope = ['https://spreadsheets.google.com/feeds',
                 'https://www.googleapis.com/auth/drive']
        creds = ServiceAccountCredentials.from_json_keyfile_name('./client_secret2.json', scope)
        client = gspread.authorize(creds)
        sheet = client.open('Hours_observation_TS_TN').sheet1
        # Extract and print all of the values
        # list_of_hashes = sheet.get_all_records()
        col2 = sheet.col_values(2)
        col5 = sheet.col_values(6)
        col5 = col5[13:]
        col5 = [float(col5[i].replace(",", ".")) for i in range(0, len(col5))]
        target_observed_tstn = pd.Series(col2[13:])
        hours_observed_tstn = pd.Series(col5)
        telescopes = ['TRAPPIST'] * len(target_observed_tstn)
        df_google_doc = pd.DataFrame({'Target': target_observed_tstn, ' Observation Hours ': hours_observed_tstn,
                                      'telescope': telescopes})
        df_google_doc.to_csv(path_spock + '/survey_hours/ObservationHours_TRAPPIST.txt', sep=',', index=False)
        subprocess.Popen(["sshpass", "-p", pwd_appcs, "scp", path_spock + '/survey_hours/ObservationHours_TRAPPIST.txt',
                          'speculoos@appcs.ra.phy.cam.ac.uk:/appct/data/SPECULOOSPipeline/spock_files/survey_hours/'])

    def update_nb_hours_sno(self):
        get_hours_files_sno()
        target_list = pd.read_csv(self.target_list, delimiter=',')
        df = pd.read_csv(path_spock + '/survey_hours/ObservationHours.txt', delimiter=',')
        sno_in_targetlist, targetlist_in_sno = index_list1_list2(df['Target'], target_list['Sp_ID'])
        target_list['nb_hours_surved'][sno_in_targetlist] = df[' Observation Hours '][targetlist_in_sno]
        # target_list.to_csv(self.target_list, sep=' ', index=False)

    def update_telescope_sno(self):
        try:
            get_hours_files_sno()
        except TimeoutError:
            print(Fore.RED + 'ERROR:  ' + Fore.BLACK + ' Are on the Liege  VPN ?')
        target_list = pd.read_csv(self.target_list, delimiter=',')
        df = pd.read_csv(path_spock + '/survey_hours/ObservationHours.txt', delimiter=',')
        sno_in_targetlist, targetlist_in_sno = index_list1_list2(df['Target'], target_list['Sp_ID'])
        df_tel = ['Artemis'] * len(df['Target'][targetlist_in_sno])
        target_list['telescope'][sno_in_targetlist] = df_tel
        target_list.to_csv(self.target_list, sep=' ', index=False)

    def update_nb_hours_from_server(self, user='educrot', password=pwd_portal):
        TargetURL = "http://www.mrao.cam.ac.uk/SPECULOOS/reports/SurveyTotal"
        target_list = pd.read_csv(self.target_list, delimiter=',')
        resp = requests.get(TargetURL, auth=(user, password))
        content = resp.text.replace("\n", "")
        open(path_spock + '/survey_hours/SurveyTotal.txt', 'wb').write(resp.content)
        df = pd.read_csv(path_spock + '/survey_hours/SurveyTotal.txt', delimiter=' ', skipinitialspace=True,
                         error_bad_lines=False)
        df['Target'][9] = 'Sp0004-2058'
        df['Target'] = [x.strip().replace('SP', 'Sp') for x in df['Target']]
        server_in_targetlist, targetlist_in_server = index_list1_list2(df['Target'], target_list['Sp_ID'])
        target_list['nb_hours_surved'][server_in_targetlist] = df['Hours'][targetlist_in_server]
        # target_list.to_csv(self.target_list, sep=' ', index=False)

    def update_telescope_from_server(self, user='educrot', password=pwd_portal):
        TargetURL = "http://www.mrao.cam.ac.uk/SPECULOOS/reports/SurveyByTelescope"
        target_list = pd.read_csv(self.target_list, delimiter=',')
        resp = requests.get(TargetURL, auth=(user, password))
        content = resp.text.replace("\n", "")
        open(path_spock + '/survey_hours/SurveyByTelescope.txt', 'wb').write(resp.content)
        df = pd.read_csv(path_spock + '/survey_hours/SurveyByTelescope.txt', delimiter=' ', skipinitialspace=True,
                         error_bad_lines=False)
        df = df.sort_values(['Target'])
        df.to_csv(path_spock + '/survey_hours/SurveyByTelescope.txt', sep=' ', index=False)
        
    def time_of_night(self, day):
        '''

        :param day: day str format '%y%m%d HH:MM:SS.sss'
        :return:
        '''
        dt_1day = Time('2018-01-02 00:00:00', scale='tt') - Time('2018-01-01 00:00:00', scale='tt')
        if self.telescope == 'Io' or self.telescope == 'Europa' or self.telescope == 'Ganymede' \
                        or self.telescope == 'Callisto':
            self.start_night = Time((Time(self.observatory.twilight_evening_nautical(day,
                                                       which='next')).value +
                                             Time(self.observatory.twilight_evening_civil(
                                                 day,
                                                 which='next')).value) / 2,
                                            format='jd')

            self.end_night = Time((Time(self.observatory.twilight_morning_nautical(day + 1,
                                                       which='nearest')).value +
                                           Time(self.observatory.twilight_morning_civil(
                                               day + 1,
                                               which='nearest')).value) / 2,
                                          format='jd')
        if self.telescope == "Artemis":
            self.start_night = Time(self.observatory.sun_set_time(day, which='next',
                                                                              horizon=-12 * u.degree).iso)
            self.end_night = Time(self.observatory.sun_rise_time(day, which='next',
                                                                             horizon=-12 * u.degree).iso)
        if self.telescope == "Saint-Ex":
            self.start_night = Time(self.observatory.sun_set_time(day, which='next',
                                                            horizon=-8.19 * u.degree).iso)
            self.end_night = Time(self.observatory.sun_rise_time(day, which='next',
                                                           horizon=-8.19 * u.degree).iso)
            if (self.start_night > self.end_night) or (self.end_night - self.start_night).value > 1:
                sys.exit('ERROR: Problem with start/end of night  on Saint-Ex !!')
        # if self.telescope == 'Saint-Ex':
        #     dura = end_saint_ex - start_saint_ex
        # else:
        self.night_duration = self.end_night - self.start_night
        # dura = Time(Time(self.observatory.twilight_morning_nautical(day + dt_1day ,which='nearest')).jd - \
        # Time(self.observatory.twilight_evening_nautical(day ,which='next')).jd,format='jd')
        # return dura

    def make_schedule(self, Altitude_constraint=None, Moon_constraint=None):
        self.Altitude_constraint = Altitude_constraint
        self.Moon_constraint = Moon_constraint
        start = time.time()
        # for telescope in self.telescopes:
        # self.nb_hours = np.zeros((len(self.target_table_spc['nb_hours_surved']),
        #                           len(self.target_table_spc['nb_hours_surved'])))
        # self.nb_hours[0, :] = self.target_table_spc['nb_hours_surved']
        # self.nb_hours[1, :] = self.target_table_spc['nb_hours_surved']
        self.targets = target_list_good_coord_format(self.target_list)

        ## This is to check whether we have target duplications on the telescopes on a same night. 
        observed_targets_SSO = []
        observed_targets_SNO = []
        observed_targets_SaintEx = []

        for i in tqdm(range(self.date_range_in_days), desc="Updating hours of obs "):
            pass
            date = self.date_range[0] + i
            day_fmt = Time(date.iso, out_subfmt='date').iso
            observed_targets_SSO += sso_planned_targets(day_fmt, self.telescope)
            observed_targets_SNO += sno_planned_targets(day_fmt)
            observed_targets_SaintEx += saintex_planned_targets(day_fmt)

        self.idx_planned_sso, self.idx_SSO_in_planned = index_list1_list2(observed_targets_SSO,
                                                                          self.target_table_spc['Sp_ID'])
        self.idx_planned_sno, self.idx_sno_in_planned = index_list1_list2(observed_targets_SNO,
                                                                          self.target_table_spc['Sp_ID'])
        self.idx_planned_saintex, self.idx_saintex_in_planned = index_list1_list2(observed_targets_SaintEx,
                                                                          self.target_table_spc['Sp_ID'])

        if self.Altitude_constraint:
            self.constraints.append(AltitudeConstraint(min=float(self.Altitude_constraint) * u.deg))
        if self.Moon_constraint:
            self.constraints.append(MoonSeparationConstraint(min=float(self.Moon_constraint) * u.deg))

        if str(self.strategy) == 'continuous':
            self.reverse_df1 = reverse_observability(self.observatory, self.targets, self.constraints, self.time_ranges)

            for t in tqdm(range(0, self.date_range_in_days), desc="Scheduling "):
                pass
                print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' day is : ', Time(self.date_range[0] + t).iso)
                self.day = self.date_ranges_day_by_day[t]
                self.time_of_night(self.day)
                # horizon_for_set_and_rise = -12 *u.degree
                # self.start_night = self.observatory.sun_set_time(self.day, which='next', horizon=horizon_for_set_and_rise)
                # self.end_night = self.observatory.sun_rise_time(self.day, which='next', horizon=horizon_for_set_and_rise)
                self.table_priority_prio(self.day)
                 self.idx_nightly_targets(t=1)
                #self.update_hours(self.day,targets)
                # if self.is_constraints_met_first_target(t):
                #     self.first_target = self.priority[self.idx_first_target]
                #     self.first_target_by_day.append(self.first_target)
                #     self.idx_first_target_by_day.append(self.idx_first_target)
                # self.update_hours_observed_first(day)

                # if not self.is_constraints_met_first_target(t):
                #     print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK +
                #           ' impossible to find the first target that respects the constraints')

                # if (self.idx_second_target is not None) and self.is_constraints_met_second_target(t):
                #     self.second_target = self.priority[self.idx_second_target]
                #     self.second_target_by_day.append(self.second_target)
                #     self.idx_second_target_by_day.append(self.idx_second_target)
                # self.update_hours_observed_second(day)

                if self.idx_second_target is None:
                    print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK + ' no second target')

                self.night_block = self.schedule_blocks(self.day)
                if self.read_locked_target:
                    import SPOCK.short_term_scheduler as SPOCKST
                    schedule_short = SPOCKST.Schedules()
                    schedule_short.load_parameters()
                    schedule_short.day_of_night = self.day
                    schedule_short.observatory = self.observatory
                    schedule_short.telescope = self.telescope
                    try:
                        schedule_short.SS1_night_blocks = Table.read(path_spock + '/DATABASE/' +
                                                                     schedule_short.telescope + '/' + 'Locked_obs/' +
                                                                     'lock_night_block_' + schedule_short.telescope +
                                                                     '_' +
                                                                     Time(schedule_short.day_of_night.iso,
                                                                          out_subfmt='date').iso + '.txt',
                                                                     format='ascii')
                        schedule_short.scheduled_table = self.night_block
                        schedule_short.planification()
                        self.night_block = schedule_short.scheduled_table_sorted
                    except FileNotFoundError:
                        pass

                self.night_block_by_day.append(self.night_block)
                self.make_night_block(self.day)
                #self.make_night_block(day)

        if str(self.strategy) == 'segmented':
            print()

    def idx_nightly_targets(self, t=1):
        """
            Give the index of the first and second targets as well as the
            corresponding row in the priority table

        Parameters
        ----------
            t: int
               number of days from start of schedule 
               We now work on a night by night basis so I put t = 1.

        Returns
        -------
            idx_first_target: int
                index first target in target list
            first_target: int
                row number idx_first_target in priority table
            idx_second_target: int
                index second target in target list
            second_target: row number idx_second_target in priority table

        """
        #idx_init_first = -1
        #self.idx_first_target = self.index_prio[idx_init_first]

        self.first_target = self.priority_ranked[np.argmax(self.priority_ranked['priority'])]  #self.priority[self.idx_first_target]

        for j in np.arange(len(self.priority_ranked)):
            nb_hours_df = self.update_hours(self.day, self.first_target)
            if nb_hours_df['nb_hours_observed'] + nb_hours_df['nb_hours_planned']>200:
                self.first_target = self.priority_ranked[np.argmax(self.priority_ranked['priority'])-j]
                continue
            else:
                break
        
        #print(nb_hours_df)
        self.selected_first_target = [t for t in self.targets if t.name == self.first_target['target_name']]
        self.idx_first_target =  np.where(self.priority_ranked['target_name'] == self.first_target['target_name'])[0]

        #dt_1day = Time('2018-01-02 00:00:00', scale='tcg') - Time('2018-01-01 00:00:00', scale='tcg')  # 1 day

        for i in range(self.idx_first_target[0] - 1, 0, -1):
            # print(self.targets[self.idx_first_target].name)
            self.second_target = self.priority_ranked[i]  #self.priority[self.idx_first_target]
            self.idx_second_target =  np.where(self.priority_ranked['target_name'] == self.second_target['target_name'])[0]
            nb_hours_df = self.update_hours(self.day, self.second_target)
            if nb_hours_df['nb_hours_observed'] + nb_hours_df['nb_hours_planned']>200:
                self.second_target = self.priority_ranked[np.argmax(self.priority_ranked['priority'])-j]
                self.selected_second_target = [t for t in self.targets if t.name == self.second_target['target_name']]
                rise_first_target = self.observatory.target_rise_time(self.date_range[0] + t,
                                                                        self.selected_first_target,
                                                                        which='next',
                                                                        horizon=self.Altitude_constraint * u.deg)
                rise_second_target = self.observatory.target_rise_time(self.date_range[0] + t,
                                                                self.selected_second_target,
                                                                which='next',
                                                                horizon=self.Altitude_constraint * u.deg)
                set_first_target = self.observatory.target_set_time(self.date_range[0] + t,
                                                                    self.selected_first_target,
                                                                    which='next',
                                                                    horizon=self.Altitude_constraint * u.deg)
                
                set_second_target = self.observatory.target_set_time(self.date_range[0] + t,
                                                                self.selected_second_target,
                                                                which='next',
                                                                horizon=self.Altitude_constraint * u.deg)

                if set_first_target < self.observatory.target_rise_time(self.date_range[0] + t,
                                                            self.selected_first_target,
                                                            which='nearest',
                                                            horizon=self.Altitude_constraint * u.deg):  # typically happens for Saint-Ex as the night start in the middle of a UTC time night

                    set_first_target = self.observatory.target_set_time(self.date_range[0] + t + 1,
                                                                        self.selected_first_target,
                                                                        which='next',
                                                                        horizon=self.Altitude_constraint * u.deg)

                if rise_first_target  < self.observatory.target_rise_time(self.date_range[0] + t,
                                                            self.selected_first_target,
                                                            which='nearest',
                                                            horizon=self.Altitude_constraint * u.deg): # typically happens for Saint-Ex as the night start in the middle of a UTC time night
                    
                    rise_first_target = self.observatory.target_rise_time(self.date_range[0] + t + 1,
                                                                            self.selected_first_target,
                                                                            which='next',
                                                                            horizon=self.Altitude_constraint * u.deg) 

                if self.telescope == "Saint-Ex":
                    if self.first_target['both']:
                        # self.idx_second_target = self.idx_first_target
                        self.second_target = self.first_target
                        break
                        
                    if self.first_target['rise'] :
                        if self.second_target['set'].any() or self.first_target['both'].any():
                            if (rise_second_target < self.start_night) and \
                                    (set_second_target > rise_first_target):
                                break

                            else:
                                self.second_target = None

                    if self.first_target['set']:
                        if self.second_target['rise'].any() or self.first_target['both'].any() :
                            if (set_second_target > self.end_night) and \
                                    (rise_second_target < set_first_target):
                                break

                            else:
                                self.second_target = None


                if (self.telescope == 'Io') or (self.telescope == 'Europa') or (self.telescope == 'Ganymede') or \
                    (self.telescope == 'Callisto') or (self.telescope == 'Artemis'):

                    if self.first_target['rise'] :
                        if self.second_target['set'].any() or self.first_target['both'].any():
                            if (rise_second_target < self.start_night) and \
                                    (set_second_target > rise_first_target):
                                break

                            else:
                                self.second_target = None

                    if self.first_target['set']:
                        if self.priority['rise'].any() or self.priority['both'].any() :
                            if (set_second_target > self.end_night) and \
                                    (rise_second_target < set_first_target):
                                print(i,"set, break")
                                break

                            else:
                                self.second_target = None
                                #self.idx_second_target = None

                    if self.first_target['both']:
                        self.second_target = self.first_target
                        break
                if self.second_target is None:
                    print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' no second target available')
            else:
                self.selected_second_target = [t for t in self.targets if t.name == self.second_target['target_name']]
                rise_first_target = self.observatory.target_rise_time(self.date_range[0] + t,
                                                                    self.selected_first_target,
                                                                    which='next',
                                                                    horizon=self.Altitude_constraint * u.deg)
                rise_second_target = self.observatory.target_rise_time(self.date_range[0] + t,
                                                                self.selected_second_target,
                                                                which='next',
                                                                horizon=self.Altitude_constraint * u.deg)
                set_first_target = self.observatory.target_set_time(self.date_range[0] + t,
                                                                    self.selected_first_target,
                                                                    which='next',
                                                                    horizon=self.Altitude_constraint * u.deg)
                
                set_second_target = self.observatory.target_set_time(self.date_range[0] + t,
                                                                self.selected_second_target,
                                                                which='next',
                                                                horizon=self.Altitude_constraint * u.deg)

                if set_first_target < self.observatory.target_rise_time(self.date_range[0] + t,
                                                            self.selected_first_target,
                                                            which='nearest',
                                                            horizon=self.Altitude_constraint * u.deg):  # typically happens for Saint-Ex as the night start in the middle of a UTC time night

                    set_first_target = self.observatory.target_set_time(self.date_range[0] + t + 1,
                                                                        self.selected_first_target,
                                                                        which='next',
                                                                        horizon=self.Altitude_constraint * u.deg)

                if rise_first_target  < self.observatory.target_rise_time(self.date_range[0] + t,
                                                            self.selected_first_target,
                                                            which='nearest',
                                                            horizon=self.Altitude_constraint * u.deg): # typically happens for Saint-Ex as the night start in the middle of a UTC time night
                    
                    rise_first_target = self.observatory.target_rise_time(self.date_range[0] + t + 1,
                                                                            self.selected_first_target,
                                                                            which='next',
                                                                            horizon=self.Altitude_constraint * u.deg) 

                if self.telescope == "Saint-Ex":
                    if self.first_target['both']:
                        # self.idx_second_target = self.idx_first_target
                        self.second_target = self.first_target
                        break
                        
                    if self.first_target['rise'] :
                        if self.second_target['set'].any() or self.first_target['both'].any():
                            if (rise_second_target < self.start_night) and \
                                    (set_second_target > rise_first_target):
                                break

                            else:
                                self.second_target = None

                    if self.first_target['set']:
                        if self.second_target['rise'].any() or self.first_target['both'].any() :
                            if (set_second_target > self.end_night) and \
                                    (rise_second_target < set_first_target):
                                break

                            else:
                                self.second_target = None


                if (self.telescope == 'Io') or (self.telescope == 'Europa') or (self.telescope == 'Ganymede') or \
                    (self.telescope == 'Callisto') or (self.telescope == 'Artemis'):

                    if self.first_target['rise'] :
                        if self.second_target['set'].any() or self.first_target['both'].any():
                            if (rise_second_target < self.start_night) and \
                                    (set_second_target > rise_first_target):
                                break

                            else:
                                self.second_target = None

                    if self.first_target['set']:
                        if self.priority['rise'].any() or self.priority['both'].any() :
                            if (set_second_target > self.end_night) and \
                                    (rise_second_target < set_first_target):
                                print(i,"set, break")
                                break

                            else:
                                self.second_target = None
                                #self.idx_second_target = None

                    if self.first_target['both']:
                        self.second_target = self.first_target
                        break
                    
                if self.second_target is None:
                    print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' no second target available')

            
    #   return [self.first_target,self.second_target]

    def table_priority_prio(self, day):

        self.priority =  pd.DataFrame(columns=['priority', 'target_name', 'set', 'rise', 'both', 'moon', 'program',
                                               'SNR_ANDOR', 'SNR_SPIRIT', 'boost', 'to_do', 'started', 'completed','SPIRIT_observable'
                           ])
        try:
            self.init_priority_table(self.day)  # initialise priority table
        except ValueError:
            sys.exit(Fore.RED + 'ERROR:  ' + Fore.BLACK + ' This is a known error, please re-run')

        idx_on_going = (self.target_table_spc['nb_hours_surved'] > 0) & (self.target_table_spc['nb_hours_surved'] < 200)
        idx_to_be_done = self.target_table_spc['nb_hours_surved'] == 0
        idx_done = self.target_table_spc['nb_hours_surved'] > 200
        
        self.priority['started'][idx_on_going] = True
        self.priority['to_do'][idx_to_be_done] = True
        self.priority['completed'][idx_done] = True

        # set priority in terms of Program number
        if self.telescope == "Callisto":
            self.priority['priority'][self.priority['Program'] == 1] = self.priority['SNR_SPIRIT'][self.priority['Program'] == 1] * 10
            self.priority['priority'][self.priority['Program'] == 3] = self.priority['SNR_SPIRIT'][self.priority['Program'] == 3] * 5

        else:
            self.priority['priority'][self.priority['Program'] == 1] = self.priority['SNR_ANDOR'][self.priority['Program'] == 1] * 10
            self.priority['priority'][self.priority['Program'] == 3] = self.priority['SNR_ANDOR'][self.priority['Program'] == 3] * 5

        # define if target is rising or setting or both during the night
        nb_reso_grid = 30
        horizon_for_set_and_rise = -12 * u.degree # degrees

        if self.telescope == 'Saint-Ex':
            delta_midnight_start = np.linspace(0, self.observatory.sun_rise_time(day, which='next',
                                                                                    horizon=-8.19 * u.degree).jd -
                                                self.observatory.sun_set_time(day, which='nearest',
                                                                                horizon=-8.19 * u.degree).jd,
                                                nb_reso_grid) * u.day  # Delta at the first day of schedule
            frame_start = AltAz(obstime=self.start_night + delta_midnight_start,
                                location=self.observatory.location)

            delta_midnight_end = np.linspace(0, self.observatory.sun_rise_time(day + self.date_range_in_days,
                                                                                which='next',
                                                                                horizon=-8.19 * u.degree).jd
                                                - self.observatory.sun_set_time(day + self.date_range_in_days,
                                                                                which='nearest',
                                                                                horizon=-8.19 * u.degree).jd,
                                                nb_reso_grid) * u.day  # Delta at the first day of schedule
            frame_end = AltAz(obstime=self.end_night + delta_midnight_end,
                                location=self.observatory.location)

        else:
            start_night_start = Time(self.observatory.sun_set_time(day, which='next',
                                                                        horizon=horizon_for_set_and_rise).iso)
            delta_midnight_start = np.linspace(0, self.observatory.sun_rise_time(day, which='next',
                                                                                        horizon=horizon_for_set_and_rise).jd -
                                                    self.observatory.sun_set_time(day, which='nearest',
                                                                                    horizon=horizon_for_set_and_rise).jd,
                                                    nb_reso_grid) * u.day  # Delta at the first day of schedule
            frame_start = AltAz(obstime=start_night_start + delta_midnight_start,
                                location=self.observatory.location)

            start_night_end = Time(self.observatory.sun_set_time(day + self.date_range_in_days,
                                                                        which='nearest', horizon=horizon_for_set_and_rise).iso)
            delta_midnight_end = np.linspace(0, self.observatory.sun_rise_time(day + self.date_range_in_days,
                                                                                    which='next',
                                                                                    horizon=horizon_for_set_and_rise).jd
                                                    - self.observatory.sun_set_time(day + self.date_range_in_days,
                                                                                    which='nearest',
                                                                                    horizon=horizon_for_set_and_rise).jd,
                                                    nb_reso_grid) * u.day  # Delta at the first day of schedule
            frame_end = AltAz(obstime=start_night_end + delta_midnight_end, location=self.observatory.location)

        target_alt = [[target.coord.transform_to(frame_start).alt,
                        target.coord.transform_to(frame_end).alt] for target in self.targets]  # This line takes time
        target_alt_start = np.asarray(target_alt)[:, 0, :]
        target_alt_end = np.asarray(target_alt)[:, 1, :]
        
        alt_set_start = list(map(first_elem_list, target_alt_start[:]))
        alt_rise_start = list(map(last_elem_list, target_alt_start[:]))
        alt_set_end = list(map(first_elem_list, target_alt_end[:]))
        alt_rise_end = list(map(last_elem_list, target_alt_end[:]))


        df_altitudes = pd.DataFrame({ 'alt set start': alt_set_start,
                           'alt rise start': alt_rise_start, 'alt set end': alt_set_end,
                           'alt rise end': alt_rise_end, 
                           'Sp_ID': self.target_table_spc['Sp_ID']})
            

        set_targets_index = (df_altitudes['alt set start'] > self.Altitude_constraint) & \
                            (df_altitudes['alt set end'] > self.Altitude_constraint)
        self.priority['set'][set_targets_index] = True

        rise_targets_index = (df_altitudes['alt rise start'] > self.Altitude_constraint) \
                             & (df_altitudes['alt rise end'] > self.Altitude_constraint)
        self.priority['rise'][rise_targets_index] = True

        both_targets_index = (df_altitudes['alt rise start'] > self.Altitude_constraint) & \
                             (df_altitudes['alt set start'] > self.Altitude_constraint)
        self.priority['both'][both_targets_index] = True

        self.priority['priority'][both_targets_index] = self.priority['priority'][both_targets_index] * 2

        if self.observatory.name == 'SSO':
            self.priority['priority'][self.idx_planned_sso] = 0
            self.priority['priority'][self.idx_planned_sno] = 0
            self.priority['priority'][self.idx_planned_saintex] = 0

        if (self.telescope == 'Artemis') or (self.telescope == 'Saint-Ex'):
            self.priority['priority'][self.idx_planned_sso] = 0
        
        ## Add moon constraint 
        self.moon_and_visibility_constraint_table = self.is_moon_and_visibility_constraint(day)
        self.priority['moon'] = self.moon_and_visibility_constraint_table['ever observable']
        
        ## Add SPIRIT observable fields constraint
        if self.telescope == "Callisto":
            df_observable_fields_SPIRIT = pd.read_csv(path_spock + "/target_lists/observable_fields_SPIRIT.csv", sep=',')
            # Convert the relevant columns to a dictionary for easy lookup
            comparison_dict = df_observable_fields_SPIRIT.set_index('Sp_ID')['enough_comparisons'].to_dict()

            # Add the new column to your Astropy Table by mapping target_name to enough_comparisons
            self.priority['enough_comparisons'] = [comparison_dict.get(target, None) for target in self.priority['target_name']
                                                   ]
            self.priority['enough_comparisons'] = [False if x is None else x for x in self.priority['enough_comparisons']]
        
        ## Save the priority dataframe        
        df_priority = self.priority.to_pandas()
        day_fmt = Time(day.iso, out_subfmt='date').iso
        df_priority.to_csv(path_spock + '/SPOCK_files/Ranking_months_' + str(self.observatory.name) +
                          '_' + str(day_fmt) + '.csv', sep=',', index=False)

        #filter all completed target or non observable due to moon 
        self.priority = self.priority[(self.priority['completed'] != True) & (self.priority['moon'] == True)] #moon has to be set to TRUE for the target to be observable
        if self.telescope == "Callisto":
            self.priority = self.priority[(self.priority['enough_comparisons']==True)]
        # Make sure the selected targets are observable that day 
        self.priority = self.priority[(self.priority['rise'] == True) | (self.priority['set'] == True) | (self.priority['both'] == True)]

        self.priority_ranked = self.priority.copy()
        self.priority_ranked.sort('priority')
        #print()

    def init_priority_table(self, day):
        day_fmt = Time(day.iso, out_subfmt='date').iso
        #self.time_of_night(day)
        if os.path.exists(path_spock + '/SPOCK_files/Ranking_months_' + str(self.observatory.name)  + '_' + str(day_fmt)   + '.csv'):
            name_file = path_spock + '/SPOCK_files/Ranking_months_' + str(self.observatory.name) + '_' +  str(day_fmt) + '.csv'
            dataframe_ranking_months = pd.read_csv(name_file, delimiter=',')
            self.priority = Table.from_pandas(dataframe_ranking_months)
            print(Fore.GREEN + 'INFO: ' + Fore.BLACK + f' Using the existing priority table for this day {day_fmt} ' )
        else:
            df_priority = pd.DataFrame({'priority': np.zeros(len(self.targets)),
                                        'target_name': self.target_table_spc['Sp_ID'],
                                        'set': np.zeros(len(self.targets), dtype=bool),
                                        'rise': np.zeros(len(self.targets), dtype=bool),
                                        'both' : np.zeros(len(self.targets), dtype=bool),
                                        'moon' : np.zeros(len(self.targets), dtype=bool),
                                        'Program' : self.target_table_spc['Program'],
                                        'SNR_ANDOR' : self.target_table_spc['SNR_JWST_HZ_tr'],
                                        'SNR_SPIRIT' : self.target_table_spc['SNR_SPIRIT'],
                                        'boost' : np.zeros(len(self.targets)),
                                        'to_do' : np.zeros(len(self.targets), dtype=bool),
                                        'started' : np.zeros(len(self.targets), dtype=bool),
                                        'completed' : np.zeros(len(self.targets), dtype=bool)
                           })
            self.priority = Table.from_pandas(df_priority)

    def shift_hours_observation(self, target):

        date_format = "%Y-%m-%d %H:%M:%S.%f"

        if self.priority[self.priority['target_name']==target['target_name']]['set']:#[idx_target] == 'set':
            set_time_begin = datetime.strptime(self.observatory.target_set_time(self.date_range[0],
                                                                                [t for t in self.targets if t.name == target['target_name']],
                                                                                which='next', horizon=24 * u.deg).iso,
                                               date_format)
            set_time_end = datetime.strptime(self.observatory.target_set_time(self.date_range[1],
                                                                              [t for t in self.targets if t.name == target['target_name']],
                                                                              which='next', horizon=24 * u.deg).iso,
                                             date_format)
            shift_hours_observation = (set_time_begin.hour +
                                       set_time_begin.minute / 60 - set_time_end.hour - set_time_end.minute / 60)

        if self.priority[self.priority['target_name']==target['target_name']]['rise']:#[idx_target] == 'rise':
            rise_time_begin = datetime.strptime(self.observatory.target_rise_time(self.date_range[0],
                                                                                  [t for t in self.targets if t.name == target['target_name']],
                                                                                  which='next', horizon=24 * u.deg).iso,
                                                date_format)
            rise_time_end = datetime.strptime(self.observatory.target_rise_time(self.date_range[1],
                                                                                [t for t in self.targets if t.name == target['target_name']],
                                                                                which='next', horizon=24 * u.deg).iso,
                                              date_format)
            shift_hours_observation = (rise_time_end.hour + rise_time_end.minute / 60 - rise_time_begin.hour -
                                       rise_time_begin.minute / 60)

        if self.priority[self.priority['target_name']==target['target_name']]['both']:#[idx_target] == 'both':
            shift_hours_observation = 0

        else:
            shift_hours_observation = 0

        return shift_hours_observation  # hours

    def schedule_blocks(self, day):
        """
            schedule the lock thanks to astroplan tools

        Parameters
        ----------
            idx_first_target: int, index of the first target
            idx_second_target: int, index of the second target
        Returns
        -------
            SS1_night_blocks: astropy.table with name, start time, end time, duration ,
            coordinates (RE and DEC) and configuration (filter and exposure time) for each target of the night

        """

        dt_1day = Time('2018-01-02 00:00:00', scale='tcg') - Time('2018-01-01 00:00:00', scale='tcg')
        delta_day = (day - self.date_range[0]).value
        self.idx_first_target =  np.where(self.target_table_spc['Sp_ID'] == self.first_target['target_name'])[0]
        if self.second_target is not None:
            self.idx_second_target =  np.where(self.target_table_spc['Sp_ID'] == self.second_target['target_name'])[0]

        if self.idx_second_target is not None:
            shift = max(self.shift_hours_observation(self.idx_first_target),
                        self.shift_hours_observation(self.idx_second_target)) / 24  # days
        else:
            shift = self.shift_hours_observation(self.idx_first_target) / 24  # days

        dur_obs_both_target = self.night_duration(day).value * u.day
        dur_obs_set_target = (self.night_duration(
            day).value / 2 - shift / self.date_range_in_days) * u.day  # (self.night_duration(day)/(2*u.day))*u.day+1*((aa.value/30)*u.hour-t/(aa.value/2)*u.hour)
        dur_obs_rise_target = (self.night_duration(
            day).value / 2 + shift / self.date_range_in_days) * u.day  # (self.night_duration(day)/(2*u.day))*u.day-1*((aa.value/30)*u.hour-t/(aa.value/2)*u.hour)

        # start_between_civil_nautical = Time((Time(
        #     self.observatory.twilight_evening_nautical(day + dt_1day * 0,
        #                                                which='next')).value +
        #                                      Time(self.observatory.twilight_evening_civil(
        #                                          day + dt_1day * 0,
        #                                          which='next')).value) / 2,
        #                                     format='jd')

        # end_between_nautical_civil = Time((Time(
        #     self.observatory.twilight_morning_nautical(day + dt_1day * (0 + 1),
        #                                                which='nearest')).value +
        #                                    Time(self.observatory.twilight_morning_civil(
        #                                        day + dt_1day * (0 + 1),
        #                                        which='nearest')).value) / 2,
        #                                   format='jd')
        # if self.telescope == "Artemis":
        #     start_between_civil_nautical = Time(self.observatory.sun_set_time(day, which='next',
        #                                                                       horizon=-12 * u.degree).iso)
        #     end_between_nautical_civil = Time(self.observatory.sun_rise_time(day, which='next',
        #                                                                      horizon=-12 * u.degree).iso)

        # start_saint_ex = Time(self.observatory.sun_set_time(day, which='next',
        #                                                     horizon=-8.19 * u.degree).iso)
        # end_saint_ex = Time(self.observatory.sun_rise_time(day, which='next',
        #                                                    horizon=-8.19 * u.degree).iso)
        # if (start_saint_ex > end_saint_ex) or (end_saint_ex - start_saint_ex).value > 1:
        #     sys.exit('ERROR: Problem with start/end of night  on Saint-Ex !!')

        if self.telescope == 'Saint-Ex':
            constraints_set_target = self.constraints + [TimeConstraint(self.start_night,
                                                                        (self.start_night + dur_obs_set_target))]
            constraints_rise_target = self.constraints + [TimeConstraint((self.start_night + dur_obs_set_target),
                                                                         self.end_night)]
            constraints_all = self.constraints + [TimeConstraint(self.start_night, self.end_night)]

        else:
            constraints_set_target = self.constraints + [TimeConstraint(self.start_night,
                                                                        (self.start_night +
                                                                         dur_obs_set_target))]

            constraints_rise_target = self.constraints + [TimeConstraint((self.start_night +
                                                                          dur_obs_set_target),
                                                                         self.end_night)]

            constraints_all = self.constraints + [TimeConstraint(self.start_night,
                                                                 self.end_night)]

        blocks = []
        # if self.target_table_spc['texp_spc'][self.idx_first_target] == 0:
        if self.telescope == 'Saint-Ex':
            self.target_table_spc['texp_spc'][self.idx_first_target] = self.exposure_time(day=day,
                                                                                          i=self.idx_first_target)
        else:
            self.target_table_spc['texp_spc'][self.idx_first_target] = self.exposure_time(day=None,
                                                                                          i=self.idx_first_target)

        # if (self.idx_second_target is not None) and self.target_table_spc['texp_spc'][self.idx_second_target] == 0:
        if self.telescope == 'Saint-Ex':
            self.target_table_spc['texp_spc'][self.idx_second_target] = self.exposure_time(day=day,
                                                                                           i=self.idx_second_target)
        else:
            self.target_table_spc['texp_spc'][self.idx_second_target] = self.exposure_time(day=None,
                                                                                           i=self.idx_second_target)

        if self.first_target['set or rise'] == 'set':
            print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' First target is \'set\'')
            a = ObservingBlock(self.targets[self.idx_first_target], dur_obs_set_target, -1,
                               constraints=constraints_set_target,
                               configuration={"filt": str(self.target_table_spc['Filter_spc'][self.idx_first_target]),
                                              "texp": str(self.target_table_spc['texp_spc'][self.idx_first_target])})
            blocks.append(a)
            b = ObservingBlock(self.targets[self.idx_second_target], dur_obs_rise_target, -1,
                               constraints=constraints_rise_target,
                               configuration={"filt": str(self.target_table_spc['Filter_spc'][self.idx_second_target]),
                                              "texp": str(self.target_table_spc['texp_spc'][self.idx_second_target])})
            blocks.append(b)

        if self.first_target['set or rise'] == 'both':
            print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' First target is \'both\'')
            a = ObservingBlock(self.targets[self.idx_first_target], dur_obs_both_target, -1,
                               constraints=constraints_all,
                               configuration={"filt": str(self.target_table_spc['Filter_spc'][self.idx_first_target]),
                                              "texp": str(self.target_table_spc['texp_spc'][self.idx_first_target])})
            blocks.append(a)

        if self.first_target['set or rise'] == 'rise':
            print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' First target is \'rise\'')
            b = ObservingBlock(self.targets[self.idx_second_target], dur_obs_set_target, -1,
                               constraints=constraints_set_target,
                               configuration={"filt": str(self.target_table_spc['Filter_spc'][self.idx_second_target]),
                                              "texp": str(self.target_table_spc['texp_spc'][self.idx_second_target])})
            blocks.append(b)
            a = ObservingBlock(self.targets[self.idx_first_target], dur_obs_rise_target, -1,
                               constraints=constraints_rise_target,
                               configuration={"filt": str(self.target_table_spc['Filter_spc'][self.idx_first_target]),
                                              "texp": str(self.target_table_spc['texp_spc'][self.idx_first_target])})
            blocks.append(a)

        transitioner = Transitioner(slew_rate=11 * u.deg / u.second)
        seq_schedule_SS1 = Schedule(self.date_range[0] + dt_1day * delta_day, self.date_range[0] +
                                    dt_1day * (delta_day + 1))
        sequen_scheduler_SS1 = SPECULOOSScheduler(constraints=constraints_all, observer=self.observatory,
                                                  transitioner=transitioner)
        sequen_scheduler_SS1(blocks, seq_schedule_SS1)

        SS1_night_blocks = seq_schedule_SS1.to_table()
        name_all = SS1_night_blocks['target']
        name = []
        for i, nam in enumerate(name_all):
            name.append(nam)
        return SS1_night_blocks

    def make_night_block(self, day):

        day_fmt = Time(day.iso, out_subfmt='date').iso

        self.night_block.add_index('target')

        try:
            index_to_delete = self.night_block.loc['TransitionBlock'].index
            self.night_block.remove_row(index_to_delete)
        except KeyError:
            print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' no transition block')

        panda_table = self.night_block.to_pandas()
        panda_table.to_csv(os.path.join(
            path_spock + '/night_blocks_propositions/' + 'night_blocks_' + self.telescope + '_' + str(
                day_fmt) + '.txt'), sep=' ', index_label='target')

    # def is_constraint_hours(self, idx_target):
    #     """
    #         Check if number of hours is ok

    #     Parameters
    #     ----------
    #         idx_target: int, index of the target you want to check

    #     Returns
    #     -------
    #         is_hours_constraint_met_target: boolean, say the hour constraint is ok or not

    #     """
    #     # nb_hours_observed = self.target_table_spc['nb_hours_surved']
    #     is_hours_constraint_met_target = True
    #     a = (1 - self.target_table_spc['nb_hours_surved'][idx_target] /
    #          (self.target_table_spc['nb_hours_threshold'][idx_target] + 10))
    #     if a < 5E-2:
    #         is_hours_constraint_met_target = False
    #     return is_hours_constraint_met_target

    def info_obs_possible(self, day):
        '''

        :param day: day str format '%y%m%d HH:MM:SS.sss'
        :return:
        '''
        duration_obs_possible = self.is_moon_and_visibility_constraint(day)[
            'fraction of time observable']  # * Time(self.night_duration(day)).value  #np.subtract(np.asarray(self.set_time_targets(day).value) , np.asarray(self.rise_time_targets(day).value))
        duration_obs_possible[duration_obs_possible < 0] = 0
        return duration_obs_possible

    def rise_time_targets(self, day):
        '''

        :param day: day str format '%y%m%d HH:MM:SS.sss'
        :return:
        '''
        rise_time = self.observatory.target_rise_time(day, self.targets, which='next', horizon=24 * u.deg)
        return rise_time

    def set_time_targets(self, day):
        '''

        :param day: day str format '%y%m%d HH:MM:SS.sss'
        :return:
        '''
        dt_1day = Time('2018-01-02 00:00:00', scale='tcg') - Time('2018-01-01 00:00:00', scale='tcg')
        set_time = self.observatory.target_set_time(day + 1 * dt_1day, self.targets, which='nearest',
                                                    horizon=24 * u.deg)
        return set_time

    def idx_is_julien_criterion(self, day):
        '''

        :param day: day str format '%y%m%d HH:MM:SS.sss'
        :return:
        '''
        nb_hour_wanted = 6  # hours
        idx_is_julien_criterion = (self.info_obs_possible(day) > nb_hour_wanted)
        return idx_is_julien_criterion

    def is_moon_and_visibility_constraint(self, day):
        """
            Check if number of moon not too close and target is visible

        Parameters
        ----------
            day: int, index of the target you want to check
        Returns
        -------
            is_hours_constraint_met_target: boolean, say the hour constraint is ok or not

        """
        constraints = self.constraints + [MoonSeparationConstraint(min=35 * u.deg)]

        dt_1day = Time('2018-01-02 00:00:00', scale='utc') - Time('2018-01-01 00:00:00', scale='utc')

        self.observability_table_day = observability_table(self.constraints, self.observatory, self.targets,
                                                                time_range=Time([self.start_night.iso,
                                                                                 self.end_night.iso]))

        self.observability_table_day['fraction of time observable'] = \
            self.observability_table_day['fraction of time observable'] * self.night_duration.value * 24

        is_visible_mid_night = is_observable(constraints, self.observatory, self.targets,
                                             times=Time(self.start_night +
                                                        self.night_duration.value / 2))

        idx_not_visible_mid_night = np.where((is_visible_mid_night == False))

        self.observability_table_day['ever observable'][idx_not_visible_mid_night[0]] = False

        return self.observability_table_day

    # def is_constraints_met_first_target(self, t):
    #     """

    #     Parameters
    #     ----------
    #     t

    #     Returns
    #     -------

    #     """

    #     # if t == 4:
    #     #     print()
    #     dt_1day = Time('2018-01-02 00:00:00', scale='tcg') - Time('2018-01-01 00:00:00', scale='tcg')

    #     # is_moon_constraint_met_first_target = \
    #     #     self.moon_and_visibility_constraint_table['ever observable'][self.idx_first_target]
    #     #hours_constraint_first = self.is_constraint_hours(self.idx_first_target)

    #     idx_safe = 1
    #     idx_safe_1rst_set = 1
    #     idx_safe_1rst_rise = 1
    #     moon_idx_set_target = 0
    #     moon_idx_rise_target = 0

    #     # if self.telescope == "Callisto":
    #     #     df_observable_fields_SPIRIT = pd.read_csv(path_spock + "/target_lists/observable_fields_SPIRIT.csv", sep=',')
    #     #     if np.any((df_observable_fields_SPIRIT["Sp_ID"] == self.first_target["Sp_ID"])):
    #     #         is_constraints_spirit_field_met_first_target = True
    #     #     else:
    #     #         is_constraints_spirit_field_met_first_target = False
    #     #         hours_constraint_first = False
    #     #         print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK +
    #     #         ' the second target did respect the constraints but is NOT in SPIRIT observable fields')

    #     while not (is_moon_constraint_met_first_target & hours_constraint_first):

    #         before_change_first_target = self.priority[self.idx_first_target]

    #         if before_change_first_target['set or rise'] == 'set':
    #             moon_idx_set_target += 1
    #             if moon_idx_set_target >= len(self.idx_set_targets_sorted):
    #                 idx_safe_1rst_set += 1
    #                 self.idx_first_target = self.index_prio[-idx_safe_1rst_set]
    #                 self.first_target = self.priority[self.idx_first_target]
    #                 is_moon_constraint_met_first_target = \
    #                     self.moon_and_visibility_constraint_table['ever observable'][self.idx_first_target]
    #                 hours_constraint_first = self.is_constraint_hours(self.idx_first_target)
    #             else:
    #                 self.idx_first_target = self.idx_set_targets_sorted[-moon_idx_set_target]
    #                 self.first_target = self.priority[self.idx_first_target]
    #                 if self.priority['priority'][self.idx_first_target] != float('-inf'):
    #                     is_moon_constraint_met_first_target = \
    #                         self.moon_and_visibility_constraint_table['ever observable'][self.idx_first_target]
    #                     hours_constraint_first = self.is_constraint_hours(self.idx_first_target)
    #                 else:
    #                     is_moon_constraint_met_first_target = False
    #                     hours_constraint_first = False

    #         elif before_change_first_target['set or rise'] == 'rise' and not is_moon_constraint_met_first_target:
    #             moon_idx_rise_target += 1
    #             if moon_idx_rise_target >= len(self.idx_rise_targets_sorted):
    #                 idx_safe_1rst_rise += 1
    #                 self.idx_first_target = self.index_prio[-idx_safe_1rst_rise]
    #                 self.first_target = self.priority[self.idx_first_target]
    #                 is_moon_constraint_met_first_target = \
    #                     self.moon_and_visibility_constraint_table['ever observable'][self.idx_first_target]
    #                 hours_constraint_first = self.is_constraint_hours(self.idx_first_target)
    #             else:
    #                 self.idx_first_target = self.idx_rise_targets_sorted[-moon_idx_rise_target]
    #                 self.first_target = self.priority[self.idx_first_target]
    #                 if self.priority['priority'][self.idx_first_target] != float('-inf'):
    #                     is_moon_constraint_met_first_target = \
    #                         self.moon_and_visibility_constraint_table['ever observable'][self.idx_first_target]
    #                     hours_constraint_first = self.is_constraint_hours(self.idx_first_target)
    #                 else:
    #                     is_moon_constraint_met_first_target = False
    #                     hours_constraint_first = False

    #         elif (before_change_first_target['set or rise'] != 'rise') and \
    #                 (before_change_first_target['set or rise'] != 'set') and not is_moon_constraint_met_first_target:
    #             idx_safe += 1
    #             self.idx_first_target = self.index_prio[-idx_safe]
    #             self.first_target = self.priority[self.idx_first_target]
    #             is_moon_constraint_met_first_target = \
    #                 self.moon_and_visibility_constraint_table['ever observable'][self.idx_first_target]
    #             hours_constraint_first = self.is_constraint_hours(self.idx_first_target)

                                
    #         elif self.telescope == "Callisto":
    #             df_observable_fields_SPIRIT = pd.read_csv(path_spock + "/target_lists/observable_fields_SPIRIT.csv", sep=',')
    #             if np.any((df_observable_fields_SPIRIT["Sp_ID"] == self.first_target["Sp_ID"])):
    #                 is_constraints_spirit_field_met_first_target = True
    #             else:
    #                 is_constraints_spirit_field_met_first_target = False
    #                 hours_constraint_first = False
    #                 print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK +
    #                 ' the second target did respect the constraints but is NOT in SPIRIT observable fields')


    #     if is_moon_constraint_met_first_target and hours_constraint_first:
    #         if self.telescope == "Callisto":
    #             if is_constraints_spirit_field_met_first_target:
    #                 is_constraints_met_first_target = True
    #             else:
    #                 is_constraints_met_first_target = False
    #         else:
    #             is_constraints_met_first_target = True
    #     else:
    #         is_constraints_met_first_target = False
    #     return is_constraints_met_first_target

    # def is_constraints_met_second_target(self, t):
    #     """

    #     Parameters
    #     ----------
    #     t

    #     Returns
    #     -------

    #     """

    #     dt_1day = Time('2018-01-02 00:00:00', scale='tcg') - Time('2018-01-01 00:00:00', scale='tcg')

    #     is_moon_constraint_met_second_target = self.moon_and_visibility_constraint_table['ever observable'][
    #         self.idx_second_target]
    #     hours_constraint_second = self.is_constraint_hours(self.idx_second_target)

    #     idx_safe = 1
    #     idx_safe_2nd_set = 1
    #     idx_safe_2nd_rise = 1
    #     moon_idx_set_target = 0
    #     moon_idx_rise_target = 0

    #     if self.idx_first_target == self.idx_second_target:
    #         print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' 2nd = 1ere ')
    #     if self.idx_second_target is None:
    #         print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' No second target for that night')

    #     if self.telescope == "Callisto":
    #         df_observable_fields_SPIRIT = pd.read_csv(path_spock + "/target_lists/observable_fields_SPIRIT.csv", sep=',')
    #         if np.any((df_observable_fields_SPIRIT["Sp_ID"] == self.second_target["Sp_ID"])):
    #             is_constraints_spirit_field_met_second_target = True
    #         else:
    #             is_constraints_spirit_field_met_second_target = False
    #             hours_constraint_second = False
    #             print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK +
    #             ' the second target did respect the constraints but is NOT in SPIRIT observable fields')

    #     while not (is_moon_constraint_met_second_target & hours_constraint_second):

    #         before_change_second_target = self.priority[self.idx_second_target]

    #         if before_change_second_target['set or rise'] == 'set':
    #             moon_idx_set_target += 1
    #             if moon_idx_set_target >= len(self.idx_set_targets_sorted):
    #                 idx_safe_2nd_set += 1
    #                 self.idx_second_target = self.index_prio[-idx_safe_2nd_set]
    #                 self.second_target = self.priority[self.idx_second_target]
    #                 is_moon_constraint_met_second_target = \
    #                 self.moon_and_visibility_constraint_table['ever observable'][self.idx_second_target]
    #                 hours_constraint_second = self.is_constraint_hours(self.idx_second_target)
    #             else:
    #                 self.idx_second_target = self.idx_set_targets_sorted[-(moon_idx_set_target)]
    #                 self.second_target = self.priority[self.idx_second_target]
    #                 if self.priority['priority'][self.idx_second_target] != float('-inf'):
    #                     is_moon_constraint_met_second_target = \
    #                     self.moon_and_visibility_constraint_table['ever observable'][self.idx_second_target]
    #                     hours_constraint_second = self.is_constraint_hours(self.idx_second_target)
    #                 else:
    #                     is_moon_constraint_met_second_target = False
    #                     hours_constraint_second = False

    #         elif before_change_second_target['set or rise'] == 'rise':
    #             moon_idx_rise_target += 1
    #             if moon_idx_rise_target >= len(self.idx_rise_targets_sorted):
    #                 idx_safe_2nd_rise += 1
    #                 self.idx_second_target = self.index_prio[-idx_safe_2nd_rise]
    #                 self.second_target = self.priority[self.idx_second_target]
    #                 is_moon_constraint_met_second_target = \
    #                 self.moon_and_visibility_constraint_table['ever observable'][self.idx_second_target]
    #                 hours_constraint_second = self.is_constraint_hours(self.idx_second_target)
    #             else:
    #                 self.idx_second_target = self.idx_rise_targets_sorted[-(moon_idx_rise_target)]
    #                 self.second_target = self.priority[self.idx_second_target]
    #                 if self.priority['priority'][self.idx_second_target] != float('-inf'):
    #                     is_moon_constraint_met_second_target = \
    #                     self.moon_and_visibility_constraint_table['ever observable'][self.idx_second_target]
    #                     hours_constraint_second = self.is_constraint_hours(self.idx_second_target)
    #                 else:
    #                     is_moon_constraint_met_second_target = False
    #                     hours_constraint_second = False

    #         elif (before_change_second_target['set or rise'] != 'rise') and (
    #                 before_change_second_target['set or rise'] != 'set'):
    #             if self.first_target['set or rise'] == 'rise':
    #                 moon_idx_set_target += 1
    #                 if moon_idx_set_target >= len(self.idx_set_targets_sorted):
    #                     idx_safe_2nd_set += 1
    #                     self.idx_second_target = self.index_prio[-idx_safe_2nd_set]
    #                     self.second_target = self.priority[self.idx_second_target]
    #                     is_moon_constraint_met_second_target = \
    #                     self.moon_and_visibility_constraint_table['ever observable'][self.idx_second_target]
    #                     hours_constraint_second = self.is_constraint_hours(self.idx_second_target)
    #                 else:
    #                     self.idx_second_target = self.idx_set_targets_sorted[-(moon_idx_set_target)]
    #                     self.second_target = self.priority[self.idx_second_target]
    #                     if self.priority['priority'][self.idx_second_target] != float('-inf'):
    #                         is_moon_constraint_met_second_target = \
    #                         self.moon_and_visibility_constraint_table['ever observable'][self.idx_second_target]
    #                         hours_constraint_second = self.is_constraint_hours(self.idx_second_target)
    #                     else:
    #                         is_moon_constraint_met_second_target = False
    #                         hours_constraint_second = False

    #             if self.first_target['set or rise'] == 'set':
    #                 moon_idx_rise_target += 1
    #                 if moon_idx_rise_target >= len(self.idx_rise_targets_sorted):
    #                     idx_safe_2nd_rise += 1
    #                     self.idx_second_target = self.index_prio[-idx_safe_2nd_rise]
    #                     self.second_target = self.priority[self.idx_second_target]
    #                     is_moon_constraint_met_second_target = \
    #                     self.moon_and_visibility_constraint_table['ever observable'][self.idx_second_target]
    #                     hours_constraint_second = self.is_constraint_hours(self.idx_second_target)

    #                 else:
    #                     self.idx_second_target = self.idx_rise_targets_sorted[-(moon_idx_rise_target)]
    #                     self.second_target = self.priority[self.idx_second_target]
    #                     if self.priority['priority'][self.idx_second_target] != float('-inf'):
    #                         is_moon_constraint_met_second_target = \
    #                         self.moon_and_visibility_constraint_table['ever observable'][self.idx_second_target]
    #                         hours_constraint_second = self.is_constraint_hours(self.idx_second_target)
    #                     else:
    #                         is_moon_constraint_met_second_target = False
    #                         hours_constraint_second = False

    #             if self.first_target['set or rise'] == 'both':
    #                 self.idx_second_target = self.idx_first_target
    #                 self.second_target = self.first_target
    #                 # is_moon_constraint_met_second_target = self.is_moon_and_visibility_constraint(t)
    #                 is_moon_constraint_met_second_target = \
    #                 self.moon_and_visibility_constraint_table['ever observable'][self.idx_second_target]
    #                 hours_constraint_second = self.is_constraint_hours(self.idx_second_target)
                
    #         elif self.telescope == "Callisto":
    #             df_observable_fields_SPIRIT = pd.read_csv(path_spock + "/target_lists/observable_fields_SPIRIT.csv", sep=',')
    #             if np.any((df_observable_fields_SPIRIT["Sp_ID"] == self.second_target["Sp_ID"])):
    #                 is_constraints_spirit_field_met_second_target = True
    #             else:
    #                 is_constraints_spirit_field_met_second_target = False
    #                 hours_constraint_second = False
    #                 print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK +
    #                 ' the second target did respect the constraints but is NOT in SPIRIT observable fields')


    #     if is_moon_constraint_met_second_target and hours_constraint_second:
    #         if self.telescope == "Callisto":
    #             if is_constraints_spirit_field_met_second_target:
    #                 is_constraints_met_second_target = True
    #             else:
    #                 is_constraints_met_second_target = False
    #         else:
    #             is_constraints_met_second_target = True
    #     else:
    #         is_constraints_met_second_target = False


    #     return is_constraints_met_second_target
    
    def update_hours(self,day,target):
        """
        Update the number of hours observed for the selected targets

        Args:
            day (Time object): day to consider
        targets (list of Astropy table rows): targets to consider
        """
        # if self.second_target is not None:
        #     shift = max(self.shift_hours_observation(self.first_target),
        #                 self.shift_hours_observation(self.second_target)) / 24  # days
        # else:
        #     shift = self.shift_hours_observation(self.first_target) / 24  # days
            #online

        target_name = target['target_name']
        # Today’s date
        today = Time(f"{date.today()} 15:00:00", scale='utc')

        # Loop over each day
        current_dt = today
        merged_observed_targets_this_d = []
        duration_obs_target_planned = 0

        while current_dt <= day:
            d = Time(current_dt)
            # for tel in ['Io','Europa','Callisto','Ganymede']:
            #     targets_on_sso = sso_planned_targets(date_is=d, telescope=tel)
            # targets_on_sno = sso_planned_targets(date_is=d)
            # targets_on_saintex = sso_planned_targets(date_is=d)
            # merged_observed_targets_this_d = targets_on_sso + targets_on_sno + targets_on_saintex
            # current_dt += timedelta(days=1)

            for tel in ['Io','Europa','Callisto','Ganymede', 'Artemis', 'Saint-Ex']:

                #local
                night_block_str = '/night_blocks_' + tel + '_' + d.datetime.strftime("%Y-%m-%d")  + '.txt'
                path_local = path_spock + '/DATABASE/' + tel + '/Archive_night_blocks' + night_block_str
                try:
                    nb_local = pd.read_csv(path_local, delimiter=' ', index_col=False)
                    if target_name in np.array(nb_local["target"]):
                            duration_obs_target_planned += nb_local.loc[nb_local["target"] == target_name, "duration (minutes)"]/60
                except FileNotFoundError:
                    print(Fore.YELLOW + 'WARNING: ' + Fore.BLACK + ' No plans in your local file for  ' +
                        tel + ' on the ' + d.datetime.strftime("%Y-%m-%d") + " looking online")
                    # online
                    nightb_url = "https://speculoos.withastra.io/SPECULOOS/Observations/" + tel + \
                            '/schedule/Archive_night_blocks/night_blocks_' + \
                            tel + '_' + d.datetime.strftime("%Y-%m-%d") + '.txt'
                    nightb = requests.get(nightb_url, auth=(user_portal, pwd_portal))
                    if nightb.status_code == 404:
                        pass
                    else:
                        path_online_to_local = path_spock + '/DATABASE/' + tel + '/Archive_night_blocks/night_blocks_' + tel + '_' + \
                        d.datetime.strftime("%Y-%m-%d") + '.txt'
                        open(path_online_to_local, 'wb').write(nightb.content)
                        nb_online = pd.read_csv(path_online_to_local, delimiter=' ', index_col=False)
                        if target_name in np.array(nb_online["target"]):
                            duration_obs_target_planned += nb_online.loc[nb_online["target"] == target, "duration (minutes)"]/60

            current_dt += timedelta(days=1)
        
        matches = self.target_table_spc[self.target_table_spc['Sp_ID'] == target_name]
        if len(matches) > 0:
            nb_hours_dict = {
                'nb_hours_observed': matches['nb_hours_surved'],#.values[0], # TO BE CHANGED TO AVA VALUES
                'nb_hours_planned': duration_obs_target_planned  # matches['nb_hours_planned'].values[0] # TO BE CHANGED TO ONLINE SCHEDULED VALUES
            }
        else:
            nb_hours_dict = {
                'nb_hours_observed': 0,
                'nb_hours_planned': 0
            }
        
        return nb_hours_dict

    # def update_hours_observed_first(self, day):
    #     """
    #         update number of hours observed for the corresponding first target

    #     Parameters
    #     ----------
    #         t: int, day of the month
    #         idx_first_target: int, index of the first target
    #     Returns
    #     -------
    #         self.nb_hours_observed
    #         self.nb_hours
    #         self.nb_hours[idx_first_target]

    #     """
    #     if self.idx_second_target is not None:
    #         shift = max(self.shift_hours_observation(self.idx_first_target),
    #                     self.shift_hours_observation(self.idx_second_target)) / 24  # days
    #     else:
    #         shift = self.shift_hours_observation(self.idx_first_target) / 24  # days

    #     dur_obs_both_target = self.night_duration(day).value * u.day
    #     dur_obs_set_target = (self.night_duration(day).value / 2 - shift / self.date_range_in_days) * u.day
    #     # (self.night_duration(day)/(2*u.day))*u.day+1*((aa.value/30)*u.hour-t/(aa.value/2)*u.hour)
    #     dur_obs_rise_target = (self.night_duration(day).value / 2 + shift / self.date_range_in_days) * u.day
    #     # (self.night_duration(day)/(2*u.day))*u.day-1*((aa.value/30)*u.hour-t/(aa.value/2)*u.hour)

    #     if self.first_target['set or rise'] == 'set':
    #         nb_hours__1rst_old = self.nb_hours[1, self.idx_first_target]  # hours
    #         a = dur_obs_set_target.value  # in days

    #     if self.first_target['set or rise'] == 'both':
    #         nb_hours__1rst_old = self.nb_hours[1, self.idx_first_target]  # hours
    #         a = dur_obs_both_target.value  # in days

    #     if self.first_target['set or rise'] == 'rise':
    #         nb_hours__1rst_old = self.nb_hours[1, self.idx_first_target]  # hours
    #         a = dur_obs_rise_target.value  # in days

    #     if self.first_target['set or rise'] == 'None':
    #         nb_hours__1rst_old = self.nb_hours[1, self.idx_first_target]  # hours
    #         a = dur_obs_rise_target.value  # in days

    #     self.nb_hours[0, self.idx_first_target] = nb_hours__1rst_old
    #     self.nb_hours[1, self.idx_first_target] = nb_hours__1rst_old + a * 24
    #     self.target_table_spc['nb_hours_surved'][self.idx_first_target] = nb_hours__1rst_old + a * 24

    # def update_hours_observed_second(self, day):
    #     """
    #         update number of hours observed for the corresponding second target

    #     Parameters
    #     ----------
    #         t: int, day of the month
    #         idx_second_target: int, index of the second target
    #     Returns
    #     -------
    #         self.nb_hours_observed
    #         self.nb_hours
    #         self.nb_hours[idx_second_target]

    #     """
    #     if self.idx_second_target is not None:
    #         shift = max(self.shift_hours_observation(self.idx_first_target),
    #                     self.shift_hours_observation(self.idx_second_target)) / 24  # days
    #     else:
    #         shift = self.shift_hours_observation(self.idx_first_target) / 24  # days
    #     dur_obs_both_target = self.night_duration(day).value * u.day
    #     dur_obs_set_target = (self.night_duration(day).value / 2 - shift / self.date_range_in_days) * u.day
    #     # (self.night_duration(day)/(2*u.day))*u.day+1*((aa.value/30)*u.hour-t/(aa.value/2)*u.hour)
    #     dur_obs_rise_target = (self.night_duration(day).value / 2 + shift / self.date_range_in_days) * u.day
    #     # (self.night_duration(day)/(2*u.day))*u.day-1*((aa.value/30)*u.hour-t/(aa.value/2)*u.hour)

    #     if self.second_target['set or rise'] == 'rise':
    #         nb_hours__2nd_old = self.nb_hours[1, self.idx_second_target]  # hours
    #         b = dur_obs_rise_target.value  # in days
    #     if self.second_target['set or rise'] == 'set':
    #         nb_hours__2nd_old = self.nb_hours[1, self.idx_second_target]  # hours
    #         b = dur_obs_set_target.value  # in days
    #     if self.second_target['set or rise'] == 'both':
    #         nb_hours__2nd_old = self.nb_hours[1, self.idx_second_target]  # hours
    #         b = dur_obs_both_target.value  # in days

    #     self.nb_hours[0, self.idx_second_target] = nb_hours__2nd_old
    #     self.nb_hours[1, self.idx_second_target] = nb_hours__2nd_old + b * 24
    #     self.target_table_spc['nb_hours_surved'][self.idx_second_target] = nb_hours__2nd_old + b * 24


    def reference_table(self):
        """

        Returns
        -------

        """

        a = np.zeros((len(self.date_ranges_day_by_day), len(self.target_table_spc['Sp_ID']))).astype("object")
        b = np.zeros((len(self.date_ranges_day_by_day), len(self.target_table_spc['Sp_ID']))).astype("object")
        c = np.zeros((len(self.date_ranges_day_by_day), len(self.target_table_spc['Sp_ID']))).astype("object")
        d = np.zeros((len(self.date_ranges_day_by_day), len(self.target_table_spc['Sp_ID']))).astype("object")
        idx_true = [0] * len(self.date_ranges_day_by_day)
        is_julien_criterion = [False] * len(self.date_ranges_day_by_day)

        for i, day in enumerate(self.date_ranges_day_by_day):
            a[i] = self.is_moon_and_visibility_constraint(day)['ever observable']
            b[i] = self.is_moon_and_visibility_constraint(day)['fraction of time observable']
            idx_true = len(np.where(a[i])[0])
            is_julien_criterion[i] = np.sum(self.idx_is_julien_criterion(day))
            c[i] = self.rise_time_targets(day)
            d[i] = self.set_time_targets(day)

        E = np.zeros((len(self.date_ranges_day_by_day), len(self.target_table_spc['Sp_ID']), 5)).astype("object")
        E[:, :, 0] = self.target_table_spc['Sp_ID']
        E[:, :, 1] = a
        E[:, :, 2] = b
        E[:, :, 3] = c
        E[:, :, 4] = d

        np.save('array' + str(self.observatory.name) + '_6hr' + 'prio_50_no_M6' + '.npy', E, allow_pickle=True)

        df = pd.DataFrame({'day': Time(self.date_ranges_day_by_day), 'nb_observable_target': idx_true,
                           'julien_criterion': is_julien_criterion})

        df.to_csv('nb_observable_target_prio_50_' + str(self.observatory.name) + '_6hr' + 'prio_50_no_M6' + '.csv',
                  sep=',', index=False)

    def exposure_time(self, day, i, telescope=None):
        """ calculation of the exposure time for a given target

        Parameters
        ----------
        day : date
            format 'yyyy-mm-dd'
        i : int
            index of target in target_list
        telescope : telescope name

        Returns
        -------
        float
            exposure time

        """
        if telescope is None:
            telescope = self.telescope

        #mphot computation
        Teff_target = int(self.target_table_spc['teff'][i]) # #K
        gaia_id = int(self.target_table_spc['Gaia_ID'][i]) #gaia

        if telescope == 'Callisto':
            # files used to generate SR
            efficiencyFile_SPIRIT = path_spock + '/../mphot/resources/systems/speculoos_PIRT_1280SciCam_-60.csv'
            filterFile_SPIRIT = path_spock + '/../mphot/resources/filters/zYJ.csv'
            # name to refer to the generated file
            name_SPIRIT, system_response_SPIRIT = mphot.generate_system_response(
                efficiencyFile_SPIRIT, filterFile_SPIRIT
            )
            # sky properties
            props_sky = {
                "pwv": 2.5,  # PWV [mm]
                "airmass": 1.1,  # airmass
                "seeing": 1.,  # seeing (==FWHM) ["]
            }
            props_callisto = {
                "name": name_SPIRIT,
                "plate_scale": 0.35 * (12 / 13.5),
                "N_dc": 230,
                "N_rn": 80,
                "well_depth": 55000,
                "bias_level": 0,
                "well_fill": 0.7,
                "read_time": 0.1,
                "r0": 0.5,
                "r1": 0.14,
                "ap_rad": 3
            }

            # get the precision and components used to calculate it (generates grid if not already present)
            result = mphot.get_precision_gaia(
                props_callisto, props_sky, source_id=gaia_id, Teff=Teff_target
            )
            # extract exposure time
            image_precision, binned_precision, components = result
            exposure_time = components["t [s]"]


        else:
            if (filt_ == 'z\'') or (filt_ == 'r\'') or (filt_ == 'i\'') or (filt_ == 'g\''):
                filt_ = filt_.replace('\'', '')
            if filt_ != 'I+z':
                filters = [filt_] + ['I+z', 'z', 'i', 'r']
            else:
                filters = ['I+z', 'z', 'i', 'r']
            filt_idx = 0
            filt_ = filters[filt_idx]

            efficiencyFile_ANDOR = path_spock + '/../mphot/resources/systems/speculoos_Andor_iKon-L-936_-60.csv'
            filterFile_ANDOR = path_spock + '/../mphot/resources/filters/I+z.csv'

            # name to refer to the generated file
            name_ANDOR, system_response_ANDOR = mphot.generate_system_response(
                efficiencyFile_ANDOR, filterFile_ANDOR
            )

            # sky properties
            props_sky = {
                "pwv": 2.5,  # PWV [mm]
                "airmass": 1.1,  # airmass
                "seeing": 1.,  # seeing (==FWHM) ["]
            }

            # instrument properties
            props_telescope_ANDOR = {
                "name": name_ANDOR,  # name to get SR/precision grid from file
                "plate_scale": 0.35,  # pixel plate scale ["]
                "N_dc": 0.2,  # dark current [e/pix/s]
                "N_rn": 6.328,  # read noise [e_rms/pix]
                "well_depth": 64000,  # well depth [e/pix]
                "well_fill": 0.7,  # fractional value to fill central target pixel, assuming gaussian (width function of seeing^)
                "read_time": 10.5,  # read time between images [s]
                "r0": 0.5,  # radius of telescope's primary mirror [m]
                "r1": 0.14,  # radius of telescope's secondary mirror [m]
                # "min_exp" : 0,          # optional, minimum exposure time [s]
                # "max_exp" : 120,        # optional, maximum exposure time [s]
                # "ap_rad" : 3            # optional, aperture radius [FWHM, seeing] for photometry -- 3 default == 7 sigma of Gaussian
            }

            # get the precision and components used to calculate it (generates grid if not already present)
            try:
                result = mphot.get_precision_gaia(props_telescope_ANDOR, props_sky, source_id=gaia_id, Teff=Teff_target)
            except FileNotFoundError:
                print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' Re-running the grid for mphot, can take 30s')
                result = mphot.get_precision_gaia(props_telescope_ANDOR, props_sky, source_id=gaia_id, Teff=Teff_target, override_grid=True)
            # extract exposure time
            image_precision, binned_precision, components = result
            exposure_time = components["t [s]"]

            while exposure_time < 10:

                filt_idx += 1
                filt_ = filters[filt_idx]
                efficiencyFile_ANDOR = path_spock + '/../mphot/resources/systems/speculoos_Andor_iKon-L-936_-60.csv'
                filterFile_ANDOR = path_spock + '/../mphot/resources/filters/'+str(filt_)+'.csv'
                # name to refer to the generated file
                name_ANDOR, system_response_ANDOR = mphot.generate_system_response(
                    efficiencyFile_ANDOR, filterFile_ANDOR
                )
                # instrument properties
                props_telescope_ANDOR = {
                "name": name_ANDOR,  # name to get SR/precision grid from file
                "plate_scale": 0.35,  # pixel plate scale ["]
                "N_dc": 0.2,  # dark current [e/pix/s]
                "N_rn": 6.328,  # read noise [e_rms/pix]
                "well_depth": 64000,  # well depth [e/pix]
                "well_fill": 0.7,  # fractional value to fill central target pixel, assuming gaussian (width function of seeing^)
                "read_time": 10.5,  # read time between images [s]
                "r0": 0.5,  # radius of telescope's primary mirror [m]
                "r1": 0.14,  # radius of telescope's secondary mirror [m]
                # "min_exp" : 0,          # optional, minimum exposure time [s]
                # "max_exp" : 120,        # optional, maximum exposure time [s]
                # "ap_rad" : 3            # optional, aperture radius [FWHM, seeing] for photometry -- 3 default == 7 sigma of Gaussian
                }

                # get the precision and components used to calculate it (generates grid if not already present)
                try:
                    result = mphot.get_precision_gaia(props_telescope_ANDOR, props_sky, source_id=gaia_id, Teff=Teff_target)
                except FileNotFoundError:
                    print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' Re-running the grid for mphot, can take 30s')
                    result = mphot.get_precision_gaia(props_telescope_ANDOR, props_sky, source_id=gaia_id, Teff=Teff_target, override_grid=True)
                    
                # extract exposure time
                image_precision, binned_precision, components = result                
                exposure_time = int(components["t [s]"])

        texp = exposure_time


        return texp

    def exposure_time_table(self, day):
        """ generate an exposure time table as the form of a file untitled

        Parameters
        ----------
        day : date
            date in fmt 'yyyy-mm-dd'

        Returns
        -------
        file
            file with most appropriate exposure time for each target "exposure_time_table.csv"

        """
        if day is None:
            print(Fore.GREEN + 'INFO: ' + Fore.BLACK + ' Not using moon phase in ETC')
        sso_texp = np.zeros(len(self.target_table_spc))
        sso_spirit_texp = np.zeros(len(self.target_table_spc))
        sno_texp = np.zeros(len(self.target_table_spc))
        saintex_texp = np.zeros(len(self.target_table_spc))

        for i in range(len(self.target_table_spc)):
            sso_texp[i] = self.exposure_time(day, i, 'Io')
            sso_spirit_texp[i] = self.exposure_time(day, i, 'Callisto')
            sno_texp[i] = sso_texp[i] #self.exposure_time(day, i, 'Artemis')
            saintex_texp[i] = sso_texp[i] #self.exposure_time(day, i, 'Saint-Ex')

            df = pd.DataFrame({'Sp_ID': self.target_table_spc['Sp_ID'],
                               'SSO_texp': sso_texp, 'SSO_SPIRIT_texp': sso_spirit_texp, 'SNO_texp': sno_texp,
                               'Saintex_texp': saintex_texp,  })
            df.to_csv(path_spock + '/SPOCK_files/exposure_time_table_mphot.csv', sep=',', index=False)

    def no_obs_with_different_tel(self):
        """ function to avoid observations of a similar target,
        expect for SNO and SAINT-EX were similar target observations are encouraged.

        Returns
        -------
        self.priority
            fill with 0 priority of targets already scheduled

        """
        idx_observed_SSO = self.idx_SSO_observed_targets()
        idx_observed_SNO = self.idx_SNO_observed_targets()
        idx_observed_saintex = self.idx_saintex_observed_targets()
        idx_observed_trappist = self.idx_trappist_observed_targets()
        if (self.telescope == 'TS_La_Silla'):
            # self.priority['priority'][idx_observed_SSO] = 0
            self.priority['priority'][idx_observed_SNO] = 0
            self.priority['priority'][idx_observed_saintex] = 0
        elif (self.telescope == 'TN_Oukaimeden'):
            self.priority['priority'][idx_observed_SSO] = 0
            self.priority['priority'][idx_observed_SNO] = 0
            self.priority['priority'][idx_observed_saintex] = 0
        elif (self.telescope == 'Artemis') or (self.telescope == 'Saint-Ex'):
            # self.priority['priority'][idx_observed_SSO] = 0
            self.priority['priority'][idx_observed_saintex] *= 2
            self.priority['priority'][idx_observed_SNO] *= 2
            # self.priority['priority'][idx_observed_SSO] = 0
            self.priority['priority'][idx_observed_trappist] = 0
        elif (self.telescope == 'Io') or (self.telescope == 'Europa') or \
                (self.telescope == 'Ganymede') or (self.telescope == 'Callisto'):
            # self.priority['priority'][idx_observed_trappist] = 0
            self.priority['priority'][idx_observed_saintex] = 0
            self.priority['priority'][idx_observed_SNO] = 0
            for i in range(len(idx_observed_SSO)):
                if self.target_table_spc['telescope'][idx_observed_SSO][i] != '':
                    if self.telescope not in self.target_table_spc['telescope'][idx_observed_SSO][i]:
                        self.priority['priority'][idx_observed_SSO[i]] = 0
        # print()


def read_night_block(telescope, day):
    """

    Parameters
    ----------
    telescope: name telescope
    day

    Returns
    -------

    """

    day_fmt = Time(day, scale='utc', out_subfmt='date').tt.datetime.strftime("%Y-%m-%d")
    path_local = path_spock + '/DATABASE/' + telescope + '/Archive_night_blocks/night_blocks_' + \
                 telescope + '_' + day_fmt + '.txt'

    if os.path.exists(path_local):
        scheduler_table = Table.read(path_local, format='ascii')
    else:
        nightb_url = "http://www.mrao.cam.ac.uk/SPECULOOS/Observations/" + telescope + \
                     '/schedule/Archive_night_blocks/night_blocks_' + \
                     telescope + '_' + day_fmt + '.txt'
        nightb = requests.get(nightb_url, auth=(user_portal, pwd_portal))

        if nightb.status_code == 404:
            sys.exit(Fore.RED + 'ERROR:  ' + Fore.BLACK + ' No plans on the server for this date')
        else:
            open(path_local, 'wb').write(nightb.content)
            scheduler_table = pd.read_csv(path_local, delimiter=' ',
                                          skipinitialspace=True, error_bad_lines=False)

    return scheduler_table


def make_docx_schedule(observatory_name, telescope, date_range, name_operator):
    if not os.path.exists(path_spock + '/TRAPPIST_schedules_docx'):
        os.makedirs(path_spock + '/TRAPPIST_schedules_docx')

    observatory = charge_observatories(observatory_name)[0]

    # ********* Uncomment if TRAPPIST observe SPECULOOS targets ********
    # df_speculoos = pd.read_csv(path_spock + '/target_lists/speculoos_target_list_v6.txt', delimiter=' ')
    # df_follow_up = pd.read_csv(path_spock + '/target_lists/target_transit_follow_up.txt', delimiter=' ')
    # df_special = pd.read_csv(path_spock + '/target_lists/target_list_special.txt', delimiter=' ')
    #
    # df_follow_up['nb_hours_surved'] = [0]*len(df_follow_up)
    # df_follow_up['nb_hours_threshold'] = [0] * len(df_follow_up)
    # df_special['nb_hours_surved'] = [0] * len(df_special)
    # df_special['nb_hours_threshold'] = [0] * len(df_special)
    # df_pandas = pd.DataFrame({'Sp_ID':pd.concat([df_speculoos['Sp_ID'],df_follow_up['Sp_ID'],df_special['Sp_ID']]),
    #     'RA': pd.concat([df_speculoos['RA'],df_follow_up['RA'],df_special['RA']]),
    #     'DEC': pd.concat([df_speculoos['DEC'],df_follow_up['DEC'],df_special['DEC']]),
    #     'J': pd.concat([df_speculoos['J'],df_follow_up['J'],df_special['J']]),
    #     'SpT': pd.concat([df_speculoos['SpT'], df_follow_up['SpT'], df_special['SpT']]),
    #     'nb_hours_surved': pd.concat([df_speculoos['nb_hours_surved'],df_follow_up['nb_hours_surved'],
    #                                   df_special['nb_hours_surved']]),
    #     'nb_hours_threshold': pd.concat([df_speculoos['nb_hours_threshold'],df_follow_up['nb_hours_threshold'],
    #                                      df_special['nb_hours_threshold']])})
    # df_pandas = df_pandas.drop_duplicates()
    # df = Table.from_pandas(df_pandas)

    nb_day_date_range = date_range_in_days(date_range)
    doc = Document()
    par = doc.add_paragraph()
    par_format = par.paragraph_format
    par_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par_format.space_before = Pt(0)
    par_format.space_after = Pt(6)
    run = par.add_run(observatory.name)
    run.bold = True
    font = run.font
    font.size = Pt(16)
    font.color.rgb = RGBColor(0, 0, 0)
    par = doc.add_paragraph()
    par_format = par.paragraph_format
    par_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par_format.space_before = Pt(0)
    par_format.space_after = Pt(12)
    run = par.add_run('Schedule from ' + Time(date_range[0].iso, out_subfmt='date').iso + ' to ' +
                      Time(date_range[1].iso, out_subfmt='date').iso)
    run.bold = True
    font = run.font
    font.size = Pt(16)
    font.color.rgb = RGBColor(0, 0, 0)
    par = doc.add_paragraph()
    par_format = par.paragraph_format
    par_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par_format.space_before = Pt(0)
    par_format.space_after = Pt(12)
    run = par.add_run('(Total time = 0hr, technical loss = 0hr, weather loss = 0hr, Exotime = 0hr, cometime = 0hr,   '
                      'chilean time = 0hr)')
    run.bold = True
    font = run.font
    font.size = Pt(12)
    font.color.rgb = RGBColor(255, 0, 0)
    par = doc.add_paragraph()
    par_format = par.paragraph_format
    par_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par_format.space_before = Pt(0)
    par_format.space_after = Pt(20)
    run = par.add_run(name_operator)
    run.italic = True
    font = run.font
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    par = doc.add_paragraph()
    par_format = par.paragraph_format
    par_format.space_before = Pt(16)
    par_format.space_after = Pt(0)

    for i in range(nb_day_date_range):
        day = date_range[0] + i
        # table_schedule = read_night_block(telescope, day)
        sun_set = observatory.sun_set_time(day, which='next').iso
        sun_rise = observatory.sun_rise_time(day, which='next').iso
        moon_illum = int(round(moon_illumination(day) * 100, 0)) * u.percent
        civil_twilights = [Time(observatory.twilight_evening_civil(day, which='next')).iso,
                           Time(observatory.twilight_morning_civil(day + 1, which='nearest')).iso]
        nautic_twilights = [Time(observatory.twilight_evening_nautical(day, which='next')).iso,
                            Time(observatory.twilight_morning_nautical(day + 1, which='nearest')).iso]
        astro_twilights = [Time(observatory.twilight_evening_astronomical(day, which='next')).iso,
                           Time(observatory.twilight_morning_astronomical(day + 1, which='nearest')).iso]
        start_night = nautic_twilights[0]  # table_schedule['start time (UTC)'][0]
        end_night = nautic_twilights[1]  # np.array(table_schedule['end time (UTC)'])[-1]
        night_duration = round((Time(nautic_twilights[1]) - Time(nautic_twilights[0])).jd * 24, 3) * u.hour

        run = par.add_run('Night starting on the ' + Time(day, out_subfmt='date').value)
        run.bold = True
        run.underline = True
        font = run.font
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)
        par = doc.add_paragraph()
        par_format = par.paragraph_format
        par_format.space_before = Pt(0)
        par_format.space_after = Pt(0)
        run = par.add_run('Moon illumination: ' + str(moon_illum))
        run.italic = True
        font = run.font
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)
        par = doc.add_paragraph()
        par_format = par.paragraph_format
        par_format.space_before = Pt(0)
        par_format.space_after = Pt(0)
        run = par.add_run(
            'Sunset - Sunrise: ' + '{:02d}'.format(Time(sun_set, out_subfmt='date_hm').datetime.hour) +
            'h' + '{:02d}'.format(Time(sun_set, out_subfmt='date_hm').datetime.minute) +
            '  / ' + '{:02d}'.format(Time(sun_rise, out_subfmt='date_hm').datetime.hour) + 'h' +
            '{:02d}'.format(Time(sun_rise, out_subfmt='date_hm').datetime.minute))
        run.italic = True
        font = run.font
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)
        par = doc.add_paragraph()
        par_format = par.paragraph_format
        par_format.space_before = Pt(0)
        par_format.space_after = Pt(0)
        run = par.add_run(
            'Civil/Naut./Astro. twilights: ' + \
            '{:02d}'.format(Time(civil_twilights[0], out_subfmt='date_hm').datetime.hour) + 'h' +
            '{:02d}'.format(Time(civil_twilights[0], out_subfmt='date_hm').datetime.minute) +
            '-' + '{:02d}'.format(Time(civil_twilights[1], out_subfmt='date_hm').datetime.hour) + 'h' +
            '{:02d}'.format(Time(civil_twilights[1], out_subfmt='date_hm').datetime.minute) +
            ' / ' + '{:02d}'.format(Time(nautic_twilights[0], out_subfmt='date_hm').datetime.hour) + 'h' +
            '{:02d}'.format(Time(nautic_twilights[0], out_subfmt='date_hm').datetime.minute) +
            '-' + '{:02d}'.format(Time(nautic_twilights[1], out_subfmt='date_hm').datetime.hour) + 'h' +
            '{:02d}'.format(Time(nautic_twilights[1], out_subfmt='date_hm').datetime.minute) +
            '  / ' + '{:02d}'.format(Time(astro_twilights[0], out_subfmt='date_hm').datetime.hour) + 'h' +
            '{:02d}'.format(Time(astro_twilights[0], out_subfmt='date_hm').datetime.minute) +
            '-' + '{:02d}'.format(Time(astro_twilights[1], out_subfmt='date_hm').datetime.hour) + 'h' +
            '{:02d}'.format(Time(astro_twilights[1], out_subfmt='date_hm').datetime.minute))
        run.italic = True
        font = run.font
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)
        par = doc.add_paragraph()
        par_format = par.paragraph_format
        par_format.space_before = Pt(0)
        par_format.space_after = Pt(0)
        run = par.add_run('Start-end of night (Naut. twil.): ' + '{:02d}'.format(Time(start_night).datetime.hour) +
                          'h' + '{:02d}'.format(Time(start_night).datetime.minute) +
                          ' to ' + '{:02d}'.format(Time(end_night).datetime.hour) + 'h' +
                          '{:02d}'.format(Time(end_night).datetime.minute))
        run.italic = True
        font = run.font
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)
        par = doc.add_paragraph()
        par_format = par.paragraph_format
        par_format.space_before = Pt(0)
        par_format.space_after = Pt(3)
        run = par.add_run('Night duration (Naut. twil.): ' + str(night_duration))
        run.italic = True
        font = run.font
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)

        # ******  Uncomment if you wish to make doc with existing night blocks ******
        # for j in range(len(table_schedule)):
        # if table_schedule['target'][j][-2:] == '_2':
        #     table_schedule['target'][j] = table_schedule['target'][j][:-2]
        # idx_target = np.where((df['Sp_ID'] == table_schedule['target'][j]))[0]

        #     start_time_target = table_schedule['start time (UTC)'][j]
        #     end_time_target = table_schedule['end time (UTC)'][j]
        #     config = table_schedule['configuration'][j]
        #     try:
        #         coords = SkyCoord(ra=df['RA'][idx_target].data.data[0] * u.deg, dec=df['DEC'][idx_target].data.data[0] * u.deg)
        #     except IndexError:
        #         break
        #     dist_moon = '34'
        #
        par = doc.add_paragraph()
        par_format = par.paragraph_format
        par_format.space_before = Pt(0)
        par_format.space_after = Pt(0)
        run = par.add_run(
            'From ' + '{:02d}'.format(Time(start_night, out_subfmt='date_hm').datetime.hour) + 'h' +
            '{:02d}'.format(Time(start_night, out_subfmt='date_hm').datetime.minute) + \
            ' to ' + '{:02d}'.format(Time(end_night, out_subfmt='date_hm').datetime.hour) + 'h' +
            '{:02d}'.format(Time(end_night, out_subfmt='date_hm').datetime.minute) + \
            ' : ' + 'xxxxxxx')
        run.bold = True
        font = run.font
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)
        # par = doc.add_paragraph()
        # par_format = par.paragraph_format
        # par_format.space_before = Pt(0)
        # par_format.space_after = Pt(0)
        # run = par.add_run('  Note:                                        ')
        # font = run.font
        # font.size = Pt(10)
        # font.color.rgb = RGBColor(0, 0, 0)
        par = doc.add_paragraph()
        par_format = par.paragraph_format
        par_format.space_before = Pt(0)
        par_format.space_after = Pt(0)
        # run = par.add_run(
        #     '  SPECULOOS : ' + str(df['nb_hours_surved'][idx_target].data.data[0]*u.hour) + ' of obs over ' + str(
        #         df['nb_hours_threshold'][idx_target].data.data[0]*u.hour))
        # font = run.font
        # font.size = Pt(10)
        # font.color.rgb = RGBColor(0, 0, 0)
        # par = doc.add_paragraph()
        # par_format = par.paragraph_format
        # par_format.space_before = Pt(0)
        # par_format.space_after = Pt(0)
        run = par.add_run('Jmag= ' + ' ' + ',  SpT= ' + '')  # + ', Moon at ' + str(dist_moon))
        # run = par.add_run('Jmag= ' + str(df['J'][idx_target].data.data[0]) + ',  SpT= ' + str(
        #     df['SpT'][idx_target].data[0]))  # + ', Moon at ' + str(dist_moon))
        font = run.font
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)
        par = doc.add_paragraph()
        par_format = par.paragraph_format
        par_format.space_before = Pt(0)
        par_format.space_after = Pt(3)
        run = par.add_run(' RA = ' + "" + " " + "" + " " + "" + \
                          ', DEC = ' + "" + " " + "" + " " + "" + ', ' + 'filter= ' + ', '  'texp= ')

        # run = par.add_run(' RA = ' + str('{:02d}'.format(int(coords.ra.hms[0]))) + " " +
        # str('{:02d}'.format(int(coords.ra.hms[1]))) + " " +
        # str('{:05.3f}'.format(round(coords.ra.hms[2], 3))) + \
        #     ', DEC = ' + str('{:02d}'.format(int(coords.dec.dms[0]))) + " " +
        #     str('{:02d}'.format(int(abs(coords.dec.dms[1])))) + " " +
        #     str('{:05.3f}'.format(round(abs(coords.dec.dms[2]), 3))) + ', ' + str(config[2:-2]).replace('\'',' '))
        font = run.font
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)
        par = doc.add_paragraph()
        par_format = par.paragraph_format
        par_format.space_before = Pt(16)
        par_format.space_after = Pt(0)

    font = run.font
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    if telescope == 'TS_La_Silla':
        doc.save(path_spock + '/TRAPPIST_schedules_docx/TS_' +
                 Time(date_range[0], out_subfmt='date').value.replace('-', '') + '_to_' +
                 Time(date_range[1], out_subfmt='date').value.replace('-', '') + '.docx')
    if telescope == 'TN_Oukaimeden':
        doc.save(path_spock + '/TRAPPIST_schedules_docx/TN_' +
                 Time(date_range[0], out_subfmt='date').value.replace('-', '') + '_to_' +
                 Time(date_range[1], out_subfmt='date').value.replace('-', '') + '.docx')


def date_range_in_days(date_range):
    date_format = "%Y-%m-%d %H:%M:%S.%f"
    date_start = datetime.strptime(date_range[0].value, date_format)
    date_end = datetime.strptime(date_range[1].value, date_format)
    return (date_end - date_start).days
